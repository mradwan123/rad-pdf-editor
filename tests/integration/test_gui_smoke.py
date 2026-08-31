"""Headless smoke tests for the GUI (SPEC.md Phase 1: "basic thumbnail
UI + undo/redo wired to the framework").

Runs under QT_QPA_PLATFORM=offscreen so it works without a display
server (set here, defensively, in case the environment hasn't already
- doesn't override a real display if one's configured).
"""

from __future__ import annotations

import contextlib
import os
import sys
from pathlib import Path
from typing import Any
from unittest.mock import patch

os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")

import pikepdf
import pytest
from PySide6.QtCore import QModelIndex, QPoint, QSize, Qt
from PySide6.QtGui import QCloseEvent, QGuiApplication, QKeySequence
from PySide6.QtPdf import QPdfDocument
from PySide6.QtTest import QTest
from PySide6.QtWidgets import QApplication, QDialog, QMenu, QMessageBox

from gui.controller import AppController
from gui.dialogs.bates_numbering_dialog import BatesNumberingDialog
from gui.dialogs.create_form_field_dialog import CreateFormFieldDialog
from gui.dialogs.crop_dialog import CropDialog
from gui.dialogs.docx_to_pdf_dialog import DocxToPdfDialog
from gui.dialogs.fill_form_dialog import FillFormDialog
from gui.dialogs.merge_dialog import MergeDialog
from gui.dialogs.metadata_dialog import MetadataDialog
from gui.dialogs.pdf_to_docx_dialog import PdfToDocxDialog
from gui.dialogs.pdf_to_jpg_dialog import PdfToJpgDialog
from gui.dialogs.properties_dialog import PropertiesDialog
from gui.dialogs.rotate_dialog import RotateDialog
from gui.dialogs.run_workflow_dialog import RunWorkflowDialog
from gui.dialogs.sign_dialog import SignDialog
from gui.dialogs.tab_placement_dialog import (
    PLACEMENT_NEW_TAB,
    PLACEMENT_REPLACE_CURRENT,
    TabPlacementDialog,
)
from gui.dialogs.tool_dialog_registry import TOOL_DIALOGS
from gui.dialogs.workflow_builder_dialog import WorkflowBuilderDialog
from gui.document_tab import DocumentTab
from gui.main_window import MainWindow


@pytest.fixture(scope="session")
def qapp() -> QApplication:
    # Qt allows exactly one QApplication per process - shared across
    # every test in this module/session.
    return QApplication.instance() or QApplication([])


@pytest.fixture(autouse=True)
def _isolated_app_data_dir(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv("PDFEDITOR_APP_DATA_DIR", str(tmp_path / "appdata"))


def _make_pdf(path: Path, num_pages: int) -> Path:
    pdf = pikepdf.Pdf.new()
    for _ in range(num_pages):
        pdf.add_blank_page(page_size=(300, 400))
    pdf.save(path)
    return path


def _make_colored_pdf(path: Path, num_pages: int, color: tuple[float, float, float]) -> Path:
    """Like `_make_pdf`, but with real, distinctly-colored visible
    content on every page (a filled rect, via fitz) - needed to tell a
    genuinely-rendered thumbnail apart from a blank-but-technically-
    present one by sampling actual pixel values, not just checking
    `count() > 0` (see the black-empty-tab regression tests below)."""
    import fitz

    doc = fitz.open()
    for _ in range(num_pages):
        page = doc.new_page(width=300, height=400)
        page.draw_rect(fitz.Rect(20, 20, 280, 380), color=color, fill=color)
    doc.save(str(path))
    doc.close()
    return path


def _make_damaged_pdf(path: Path, color: tuple[float, float, float]) -> Path:
    """A PDF that pikepdf reads happily but QtPdf refuses.

    Truncating the file destroys the cross-reference table and trailer.
    qpdf (behind pikepdf) reconstructs those, so `AppController.
    open_document` succeeds and reports the real page count - while
    QtPdf rejects the identical bytes with InvalidFileFormat. That
    divergence between the two engines is the whole bug being guarded
    here, so the fixture has to reproduce it rather than simulate it;
    the assertions below confirm both halves really hold.
    """
    intact = path.with_name(f"intact_{path.name}")
    _make_colored_pdf(intact, 2, color)
    data = intact.read_bytes()
    path.write_bytes(data[: int(len(data) * 0.85)])
    return path


def _thumbnail_center_pixel(
    tab: DocumentTab, window: MainWindow, index: int = 0
) -> tuple[int, int, int, int]:
    """The real rendered center-pixel color of thumbnail `index`,
    sampled from the actual built QIcon/QPixmap - not just "an item
    exists." A black/empty tab could still technically have a
    QListWidgetItem with a blank or stale icon; this catches that."""
    item = tab.thumbnail_list.item(index)
    assert item is not None, "expected a rendered thumbnail, found none"
    pixmap = item.icon().pixmap(window.thumbnail_size)
    image = pixmap.toImage()
    return image.pixelColor(image.width() // 2, image.height() // 2).getRgb()


def _open_tab(window: MainWindow, path: Path) -> DocumentTab:
    """Open `path` in a new tab through the window's own open flow.
    The New Tab / Replace Current Tab placement choice is passed
    explicitly rather than mocked, so no modal dialog is involved."""
    window._open_document_path(path, PLACEMENT_NEW_TAB)
    tab = window.current_tab
    assert tab is not None
    return tab


def _fake_placement(placement: str) -> Any:
    """Stand-in for TabPlacementDialog.exec that clicks the dialog's
    real button for `placement`, so the dialog's own handler decides
    what `placement` ends up as rather than the test setting it
    directly. TabPlacementDialog is a plain Python QDialog subclass
    precisely so patching its `exec` works at all (a QMessageBox's
    compiled `exec` silently isn't intercepted - see CLAUDE.md)."""

    def fake_exec(self: TabPlacementDialog) -> QDialog.DialogCode:
        self.button_for(placement).click()
        return QDialog.DialogCode(self.result())

    return fake_exec


def _force_close(window: MainWindow) -> None:
    """Close the window with no unsaved-changes prompt.

    A bare window.close() on a dirty document triggers a real, modal
    QMessageBox.warning() that nothing can click headlessly - it hangs
    pytest forever (CLAUDE.md documents this the hard way). Wiping
    every tab's session first is equivalent to choosing Discard on
    each, and leaves closeEvent nothing to ask about.
    """
    for tab in window.tabs():
        tab.controller.close_session()
    window.close()


def test_starts_on_empty_state_with_branded_title(qapp: QApplication) -> None:
    window = MainWindow()
    assert window.windowTitle() == "Rad PDF Editor"
    assert window.stack.currentWidget() is window.empty_state
    window.close()


def test_opening_a_document_switches_to_thumbnail_view(qapp: QApplication, tmp_path: Path) -> None:
    src = _make_pdf(tmp_path / "src.pdf", 2)
    window = MainWindow()

    tab = _open_tab(window, src)

    assert window.stack.currentWidget() is window.tab_widget
    assert window.thumbnail_list is tab.thumbnail_list
    assert window.tab_widget.count() == 1
    _force_close(window)


def test_closing_document_returns_to_empty_state(qapp: QApplication, tmp_path: Path) -> None:
    src = _make_pdf(tmp_path / "src.pdf", 1)
    window = MainWindow()
    _open_tab(window, src)

    window._close_document()

    assert window.tab_widget.count() == 0
    assert window.stack.currentWidget() is window.empty_state
    assert window.windowTitle() == "Rad PDF Editor"
    window.close()


def test_open_render_undo_redo_save_close(qapp: QApplication, tmp_path: Path) -> None:
    from core.ops.organize import RotatePagesOperation

    src = _make_pdf(tmp_path / "src.pdf", 3)
    window = MainWindow()

    assert window.thumbnail_list is None
    assert not window.undo_action.isEnabled()

    _open_tab(window, src)
    assert window.thumbnail_list.count() == 3
    assert "src.pdf" in window.windowTitle()

    window.controller.apply_operation(RotatePagesOperation(angle=90))
    window._refresh()
    assert window.undo_action.isEnabled()

    out = tmp_path / "out.pdf"
    window.controller.save_as(out)
    with pikepdf.Pdf.open(out) as pdf:
        assert len(pdf.pages) == 3
        assert int(pdf.pages[0].get("/Rotate", 0)) == 90

    window.controller.undo()
    window._refresh()
    assert not window.undo_action.isEnabled()
    assert window.redo_action.isEnabled()

    working_dir = window.controller.doc.working_path.parent
    _force_close(window)
    assert not working_dir.exists()


def test_dragging_a_thumbnail_reorders_the_document(qapp: QApplication, tmp_path: Path) -> None:
    # model().moveRow(...) triggers the exact same rowsMoved signal a
    # real mouse drag-and-drop would - InternalMove drag gestures
    # aren't reliably simulatable headlessly, but this exercises the
    # real signal-handling code path, not a hand-rolled substitute.
    src = _make_pdf(tmp_path / "src.pdf", 4)
    window = MainWindow()
    _open_tab(window, src)

    moved = window.thumbnail_list.model().moveRow(QModelIndex(), 3, QModelIndex(), 0)
    assert moved
    QTest.qWait(50)  # let the QTimer.singleShot(0, ...) deferral run

    assert len(window.controller.doc.operation_log) == 1
    applied = window.controller.doc.operation_log[-1].serialize()
    assert applied["type"] == "reorder_pages"
    assert applied["page_order"] == [4, 1, 2, 3]
    assert window.undo_action.isEnabled()
    with pikepdf.Pdf.open(window.controller.doc.working_path) as pdf:
        assert len(pdf.pages) == 4
    _force_close(window)


def test_reordering_thumbnails_only_affects_the_active_tab(
    qapp: QApplication, tmp_path: Path
) -> None:
    # Regression coverage for the multi-tab merge: _on_thumbnails_reordered
    # is connected per-tab at tab-creation time with the tab bound into
    # the lambda's default argument (`t=tab`), specifically so a
    # deferred reorder can't resolve against "whatever tab happens to
    # be active" by the time its QTimer.singleShot(0, ...) callback
    # runs. Verified here against real per-tab state, not just that the
    # signal fired once.
    a = _make_pdf(tmp_path / "a.pdf", 4)
    b = _make_pdf(tmp_path / "b.pdf", 3)
    window = MainWindow()
    tab_a = _open_tab(window, a)
    tab_b = _open_tab(window, b)
    assert window.current_tab is tab_b

    window.tab_widget.setCurrentWidget(tab_a)
    moved = tab_a.thumbnail_list.model().moveRow(QModelIndex(), 3, QModelIndex(), 0)
    assert moved
    QTest.qWait(50)

    assert [op.serialize()["type"] for op in tab_a.controller.doc.operation_log] == [
        "reorder_pages"
    ]
    assert tab_b.controller.doc.operation_log == []
    with pikepdf.Pdf.open(tab_a.controller.doc.working_path) as pdf:
        assert len(pdf.pages) == 4
    with pikepdf.Pdf.open(tab_b.controller.doc.working_path) as pdf:
        assert len(pdf.pages) == 3
    _force_close(window)


def test_tool_actions_disabled_without_open_document_except_merge(qapp: QApplication) -> None:
    window = MainWindow()
    assert not window.tool_actions["rotate_pages"].isEnabled()
    assert not window.tool_actions["watermark"].isEnabled()
    assert window.tool_actions["merge"].isEnabled()
    window.close()


def test_a_pdf_qtpdf_cannot_load_still_gets_thumbnails(
    qapp: QApplication, tmp_path: Path
) -> None:
    """Regression: "I open a PDF and there is no thumbnail."

    pikepdf repairs the damaged xref and opens the document, so a tab
    appears with the right page count - but QtPdf rejected the same
    bytes and `_render_thumbnails` logged one line and returned,
    leaving an empty grid with nothing on screen explaining why.
    """
    damaged = _make_damaged_pdf(tmp_path / "damaged.pdf", (1, 0, 0))

    # Both halves of the premise, so this can't rot into a test of a
    # file that is simply fine.
    with pikepdf.Pdf.open(damaged) as pdf:
        assert len(pdf.pages) == 2
    probe = QPdfDocument()
    assert probe.load(str(damaged)) != QPdfDocument.Error.None_

    window = MainWindow()
    tab = _open_tab(window, damaged)

    assert tab.controller.is_open
    assert tab.thumbnail_list.count() == 2
    # Rendered for real, not blank placeholders.
    red, green, blue, _alpha = _thumbnail_center_pixel(tab, window)
    assert red > 200 and green < 60 and blue < 60
    _force_close(window)


def test_an_undecodable_document_says_so_instead_of_showing_an_empty_grid(
    qapp: QApplication, tmp_path: Path
) -> None:
    """When neither engine can render, the status bar has to say so -
    an empty grid with a "3 page(s)" message underneath reads as the
    app being broken."""
    src = _make_pdf(tmp_path / "src.pdf", 2)
    window = MainWindow()
    tab = _open_tab(window, src)
    assert tab.thumbnail_list.count() == 2

    # Force both engines to fail: QtPdf rejects, then PyMuPDF raises.
    def _fail_load(self: QPdfDocument, _path: str) -> QPdfDocument.Error:
        return QPdfDocument.Error.InvalidFileFormat

    with (
        patch.object(QPdfDocument, "load", _fail_load),
        patch("gui.main_window.fitz.open", side_effect=RuntimeError("no engine")),
    ):
        window._refresh()

    assert tab.thumbnail_list.count() == 0
    assert "Could not render" in window.statusBar().currentMessage()

    # And it recovers: a later refresh with working engines renders
    # again and drops the message.
    window._refresh()
    assert tab.thumbnail_list.count() == 2
    assert "page(s)" in window.statusBar().currentMessage()
    _force_close(window)


def test_merge_from_tools_menu_opens_a_document(qapp: QApplication, tmp_path: Path) -> None:
    a = _make_pdf(tmp_path / "a.pdf", 1)
    b = _make_pdf(tmp_path / "b.pdf", 2)
    window = MainWindow()

    def fake_exec(self: MergeDialog) -> QDialog.DialogCode:
        self.file_list.addItems([str(a), str(b)])
        return QDialog.DialogCode.Accepted

    with patch.object(MergeDialog, "exec", fake_exec):
        window._run_tool("merge", MergeDialog)

    assert window.controller.is_open
    assert window.thumbnail_list.count() == 3
    assert window.tool_actions["rotate_pages"].isEnabled()
    _force_close(window)


def _make_docx(path: Path, *, body: str = "WORD BODY TEXT") -> Path:
    import docx

    document = docx.Document()
    document.add_paragraph(body)
    document.save(str(path))
    return path


def _fake_docx_dialog_exec(source: Path) -> Any:
    def fake_exec(self: DocxToPdfDialog) -> QDialog.DialogCode:
        self._source_path = source
        return QDialog.DialogCode.Accepted

    return fake_exec


def test_word_to_pdf_runs_with_no_document_open(qapp: QApplication, tmp_path: Path) -> None:
    """Regression: every external-source tool used to be gated behind
    "Open a document first" - only Merge was exempt - so Word to PDF was
    unreachable from a freshly launched window, even though the CLI has
    always accepted it with nothing open. Asserts the conversion really
    lands in a new tab, not just that no error box appeared."""
    source = _make_docx(tmp_path / "in.docx")
    window = MainWindow()
    assert window.tab_widget.count() == 0

    errors: list[str] = []
    with patch.object(
        QMessageBox, "critical", lambda *args, **kwargs: errors.append(str(args[2]))
    ), patch.object(DocxToPdfDialog, "exec", _fake_docx_dialog_exec(source)):
        window._run_tool("docx_to_pdf", DocxToPdfDialog)

    assert errors == []
    assert window.tab_widget.count() == 1
    controller = window.controller
    assert controller is not None and controller.is_open
    assert "Converted Word document to PDF" in controller.doc.operation_log[-1].describe()
    assert window.thumbnail_list.count() >= 1
    with pikepdf.Pdf.open(controller.doc.working_path) as pdf:
        assert len(pdf.pages) >= 1
    _force_close(window)


def test_external_source_tools_stay_enabled_with_no_document_open(
    qapp: QApplication,
) -> None:
    """The menu-item half of the same bug: `_update_action_state` used
    to enable only "merge" when nothing was open, so the Word to PDF
    action was greyed out and _run_tool's guard was never even
    reachable from the Tools menu."""
    window = MainWindow()
    assert window.tab_widget.count() == 0

    for tool_id in ("merge", "docx_to_pdf", "pptx_to_pdf", "xlsx_to_pdf", "html_to_pdf", "repair"):
        assert window.tool_actions[tool_id].isEnabled(), f"{tool_id} should be usable with no document open"
    # A tool that genuinely needs an open document stays disabled.
    assert not window.tool_actions["rotate_pages"].isEnabled()
    _force_close(window)


def test_word_to_pdf_with_an_unreadable_source_does_not_strand_an_empty_tab(
    qapp: QApplication, tmp_path: Path
) -> None:
    """The no-document-open path creates the tab before the operation
    runs; a source that fails to convert must not leave it behind.

    The fallback engine is forced because LibreOffice is permissive
    enough to convert a plain text file with a .docx name quite
    happily - python-docx is the engine that rejects it.
    """
    not_a_docx = tmp_path / "broken.docx"
    not_a_docx.write_text("this is plainly not a Word document")
    window = MainWindow()

    errors: list[str] = []
    with (
        patch.object(QMessageBox, "critical", lambda *args, **kwargs: errors.append(str(args[2]))),
        patch.object(DocxToPdfDialog, "exec", _fake_docx_dialog_exec(not_a_docx)),
        patch("core.ops.convert_to.libreoffice_binary", lambda: None),
    ):
        window._run_tool("docx_to_pdf", DocxToPdfDialog)

    assert errors, "a failed conversion should report an error"
    assert window.tab_widget.count() == 0
    _force_close(window)


def test_run_tool_applies_operation_via_dialog_values(qapp: QApplication, tmp_path: Path) -> None:
    src = _make_pdf(tmp_path / "src.pdf", 1)
    window = MainWindow()
    _open_tab(window, src)

    def fake_exec(self: RotateDialog) -> QDialog.DialogCode:
        self.angle.setCurrentText("180")
        return QDialog.DialogCode.Accepted

    with patch.object(RotateDialog, "exec", fake_exec):
        window._run_tool("rotate_pages", RotateDialog)

    with pikepdf.Pdf.open(window.controller.doc.working_path) as pdf:
        assert int(pdf.pages[0].get("/Rotate", 0)) == 180
    _force_close(window)


def test_cancelling_a_tool_dialog_leaves_document_unchanged(qapp: QApplication, tmp_path: Path) -> None:
    src = _make_pdf(tmp_path / "src.pdf", 1)
    window = MainWindow()
    _open_tab(window, src)
    ops_before = len(window.controller.doc.operation_log)

    def fake_cancel(self: RotateDialog) -> QDialog.DialogCode:
        return QDialog.DialogCode.Rejected

    with patch.object(RotateDialog, "exec", fake_cancel):
        window._run_tool("rotate_pages", RotateDialog)

    assert len(window.controller.doc.operation_log) == ops_before
    window.close()


def test_running_a_tool_without_open_document_shows_error_not_crash(
    qapp: QApplication,
) -> None:
    window = MainWindow()
    with patch("gui.main_window.QMessageBox.critical") as mock_critical:
        window._run_tool("rotate_pages", RotateDialog)
    mock_critical.assert_called_once()
    window.close()


def test_phase2_tools_are_all_registered_in_the_menu(qapp: QApplication) -> None:
    window = MainWindow()
    for tool_id in (
        "crop",
        "resize",
        "n_up",
        "grayscale",
        "header_footer",
        "bates_numbering",
        "flatten",
        "remove_annotations",
        "fill_form",
        "sign",
        "create_form_field",
    ):
        assert tool_id in window.tool_actions
    window.close()


def test_crop_then_bates_numbering_via_tools_menu(qapp: QApplication, tmp_path: Path) -> None:
    import pdfplumber

    src = _make_pdf(tmp_path / "src.pdf", 2)
    window = MainWindow()
    _open_tab(window, src)

    def fake_crop(self: CropDialog) -> QDialog.DialogCode:
        self.margin_top.setValue(20)
        self.margin_left.setValue(10)
        return QDialog.DialogCode.Accepted

    with patch.object(CropDialog, "exec", fake_crop):
        window._run_tool("crop", CropDialog)

    with pikepdf.Pdf.open(window.controller.doc.working_path) as pdf:
        assert [float(x) for x in pdf.pages[0].mediabox] == [10.0, 0.0, 300.0, 380.0]

    def fake_bates(self: BatesNumberingDialog) -> QDialog.DialogCode:
        self.prefix.setText("DOC-")
        return QDialog.DialogCode.Accepted

    with patch.object(BatesNumberingDialog, "exec", fake_bates):
        window._run_tool("bates_numbering", BatesNumberingDialog)

    with pdfplumber.open(window.controller.doc.working_path) as pdf:
        assert pdf.pages[0].extract_text() == "DOC-00001"
        assert pdf.pages[1].extract_text() == "DOC-00002"

    assert len(window.controller.doc.operation_log) == 2
    assert window.undo_action.isEnabled()
    _force_close(window)


def _make_pdf_with_text_field(path: Path) -> Path:
    pdf = pikepdf.Pdf.new()
    page = pdf.add_blank_page(page_size=(300, 400))
    field = pdf.make_indirect(
        pikepdf.Dictionary(
            {
                "/FT": pikepdf.Name("/Tx"),
                "/T": pikepdf.String("name"),
                "/Rect": pikepdf.Array([50, 300, 250, 320]),
                "/Subtype": pikepdf.Name("/Widget"),
                "/Type": pikepdf.Name("/Annot"),
                "/V": pikepdf.String(""),
                "/DA": pikepdf.String("/Helv 12 Tf 0 g"),
            }
        )
    )
    page.obj["/Annots"] = pikepdf.Array([field])
    pdf.Root["/AcroForm"] = pdf.make_indirect(
        pikepdf.Dictionary(
            {
                "/Fields": pikepdf.Array([field]),
                "/NeedAppearances": True,
                "/DR": pikepdf.Dictionary(
                    {
                        "/Font": pikepdf.Dictionary(
                            {
                                "/Helv": pdf.make_indirect(
                                    pikepdf.Dictionary(
                                        {
                                            "/Type": pikepdf.Name("/Font"),
                                            "/Subtype": pikepdf.Name("/Type1"),
                                            "/BaseFont": pikepdf.Name("/Helvetica"),
                                        }
                                    )
                                )
                            }
                        )
                    }
                ),
            }
        )
    )
    pdf.save(path)
    return path


def test_fill_form_via_tools_menu_uses_detected_field_names(
    qapp: QApplication, tmp_path: Path
) -> None:
    import fitz

    src = _make_pdf_with_text_field(tmp_path / "form.pdf")
    window = MainWindow()
    _open_tab(window, src)

    def fake_fill(self: FillFormDialog) -> QDialog.DialogCode:
        assert "name" in self._inputs
        self._inputs["name"].setText("Jane Smith")
        return QDialog.DialogCode.Accepted

    with patch.object(FillFormDialog, "exec", fake_fill):
        window._run_tool("fill_form", None)

    with fitz.open(window.controller.doc.working_path) as pdf:
        assert "Jane Smith" in pdf[0].get_text()
    _force_close(window)


def test_sign_via_tools_menu_places_image(qapp: QApplication, tmp_path: Path) -> None:
    from PIL import Image, ImageDraw

    src = _make_pdf(tmp_path / "src.pdf", 1)
    sig = tmp_path / "sig.png"
    img = Image.new("RGBA", (200, 80), (0, 0, 0, 0))
    ImageDraw.Draw(img).line((10, 60, 190, 20), fill=(0, 0, 200, 255), width=6)
    img.save(sig)

    window = MainWindow()
    _open_tab(window, src)

    def fake_sign(self: SignDialog) -> QDialog.DialogCode:
        self._image_path = sig
        self.page.setValue(1)
        self.x0.setValue(50)
        self.y0.setValue(50)
        self.x1.setValue(250)
        self.y1.setValue(130)
        return QDialog.DialogCode.Accepted

    with patch.object(SignDialog, "exec", fake_sign):
        window._run_tool("sign", SignDialog)

    with pikepdf.Pdf.open(window.controller.doc.working_path) as pdf:
        assert len(pdf.pages) == 1
    assert window.undo_action.isEnabled()
    _force_close(window)


def test_sign_dragged_on_the_canvas_lands_where_it_was_dropped(
    qapp: QApplication, tmp_path: Path
) -> None:
    """End-to-end for the interactive placement canvas: position the
    overlay on the rendered page by moving the real graphics item (a
    genuine mouse drag isn't simulatable offscreen - same reason
    test_dragging_a_thumbnail_reorders_the_document calls moveRow), let
    the dialog convert it, and check with fitz where the image actually
    ended up in the output PDF.

    The fixture page is 300x400 pt and the canvas renders its long edge
    to 800 px, so the scale is 2.0. An overlay at scene (100, 120)
    sized 200x80 px is therefore PDF rect (50, 300, 150, 340) measured
    from the bottom-left - which fitz, whose own Rect is top-left
    origin, reports back as bbox (50, 60, 150, 100) on a 400 pt page.
    """
    import fitz
    from PIL import Image, ImageDraw
    from PySide6.QtCore import QRectF

    src = _make_pdf(tmp_path / "src.pdf", 1)
    sig = tmp_path / "sig.png"
    # 200x80, the same 2.5 aspect ratio as the 100x40 pt target rect -
    # fitz's insert_image keeps proportion, so an image of a different
    # shape would legitimately be letterboxed inside the rect and the
    # bbox check below would be about the aspect fit, not the placement.
    img = Image.new("RGB", (200, 80), (255, 255, 255))
    ImageDraw.Draw(img).line((10, 60, 190, 20), fill=(0, 0, 200), width=6)
    img.save(sig)

    window = MainWindow()
    _open_tab(window, src)

    def fake_sign(self: SignDialog) -> QDialog.DialogCode:
        assert self.canvas is not None, "the canvas needs the open document's path"
        self.set_image_path(sig)
        item = self.canvas.placement_item()
        assert item is not None
        item.set_rect(QRectF(100, 120, 200, 80))
        self.canvas.rect_changed.emit()
        return QDialog.DialogCode.Accepted

    with patch.object(SignDialog, "exec", fake_sign):
        window._run_tool("sign", SignDialog)

    working = window.controller.doc.working_path
    assert working is not None
    with fitz.open(working) as pdf:
        images = pdf[0].get_image_info()
        assert len(images) == 1
        assert images[0]["bbox"] == pytest.approx((50.0, 60.0, 150.0, 100.0))
    _force_close(window)


@pytest.mark.parametrize("accepted", [True, False])
def test_sign_dialog_releases_the_working_file_before_the_session_is_wiped(
    qapp: QApplication, tmp_path: Path, accepted: bool
) -> None:
    """Regression: SignDialog's placement canvas previews the session
    working copy through a QPdfDocument, which holds an OS handle on it
    for as long as it is loaded, and the dialog is parented to
    MainWindow - so it outlived exec() and was still holding that
    handle when close_session() securely wiped the session dir.

    Linux/macOS unlink an open file happily, so this cost nothing
    locally; Windows refuses (WinError 32), which is how it surfaced -
    both signature tests failing in CI inside
    core/security/secure_delete.py. That makes it a real defect, not a
    test artefact: on Windows the confidential working copy was not
    being wiped on close at all.

    Asserted as "the dialog explicitly released the document", not as
    "deleting the file worked", because the latter passes on this
    platform whether the bug is fixed or not. Both the accepted and
    cancelled paths are checked - _run_tool releases in a finally.
    """
    from PIL import Image

    src = _make_pdf(tmp_path / "src.pdf", 1)
    sig = tmp_path / "sig.png"
    Image.new("RGB", (200, 80), (255, 255, 255)).save(sig)

    window = MainWindow()
    _open_tab(window, src)
    dialogs: list[SignDialog] = []

    def fake_sign(self: SignDialog) -> QDialog.DialogCode:
        dialogs.append(self)
        assert self.canvas is not None
        assert self.canvas.page_count() == 1, "the preview should be loaded here"
        self.set_image_path(sig)
        return QDialog.DialogCode.Accepted if accepted else QDialog.DialogCode.Rejected

    with patch.object(SignDialog, "exec", fake_sign):
        window._run_tool("sign", SignDialog)

    assert len(dialogs) == 1
    canvas = dialogs[0].canvas
    assert canvas is not None
    assert canvas.page_count() == 0, "the previewed document is still open"
    assert not canvas.has_page()
    _force_close(window)


def test_create_form_field_via_tools_menu_adds_a_text_field(
    qapp: QApplication, tmp_path: Path
) -> None:
    import fitz

    src = _make_pdf(tmp_path / "src.pdf", 1)
    window = MainWindow()
    _open_tab(window, src)

    def fake_create(self: CreateFormFieldDialog) -> QDialog.DialogCode:
        self.field_name.setText("full_name")
        self.field_type.setCurrentText("text")
        self.page.setValue(1)
        self.x0.setValue(50)
        self.y0.setValue(300)
        self.x1.setValue(250)
        self.y1.setValue(320)
        self.default_value.setText("Jane Doe")
        return QDialog.DialogCode.Accepted

    with patch.object(CreateFormFieldDialog, "exec", fake_create):
        window._run_tool("create_form_field", CreateFormFieldDialog)

    with fitz.open(window.controller.doc.working_path) as pdf:
        widgets = list(pdf[0].widgets())
    assert len(widgets) == 1
    assert widgets[0].field_name == "full_name"
    assert widgets[0].field_value == "Jane Doe"
    assert window.undo_action.isEnabled()
    _force_close(window)


def test_refresh_does_not_leak_qpdfdocument_instances(qapp: QApplication, tmp_path: Path) -> None:
    # Regression: _render_thumbnails used to parent its throwaway
    # QPdfDocument to `self` (MainWindow), so every _refresh() (every
    # applied operation, undo, or redo) leaked one instance for the
    # life of the window instead of being freed after rendering.
    from PySide6.QtPdf import QPdfDocument

    from core.ops.organize import RotatePagesOperation

    src = _make_pdf(tmp_path / "src.pdf", 1)
    window = MainWindow()
    _open_tab(window, src)

    before = len([c for c in window.children() if isinstance(c, QPdfDocument)])
    for _ in range(5):
        window.controller.apply_operation(RotatePagesOperation(angle=90))
        window._refresh()
    after = len([c for c in window.children() if isinstance(c, QPdfDocument)])

    assert after == before
    _force_close(window)


def test_reordering_to_the_same_order_does_not_record_a_no_op_operation(
    qapp: QApplication, tmp_path: Path
) -> None:
    # Regression: dragging (or a duplicate rowsMoved signal for the
    # same gesture) that results in the identity order previously
    # still pushed a no-op ReorderPagesOperation onto the undo stack.
    src = _make_pdf(tmp_path / "src.pdf", 3)
    window = MainWindow()
    tab = _open_tab(window, src)
    ops_before = len(window.controller.doc.operation_log)

    for i in range(tab.thumbnail_list.count()):
        tab.thumbnail_list.item(i).setData(Qt.ItemDataRole.UserRole, i + 1)
    window._apply_thumbnail_reorder(tab)

    assert len(window.controller.doc.operation_log) == ops_before
    window.close()


def test_closing_a_clean_document_does_not_prompt(qapp: QApplication, tmp_path: Path) -> None:
    src = _make_pdf(tmp_path / "src.pdf", 1)
    window = MainWindow()
    _open_tab(window, src)

    with patch.object(QMessageBox, "warning") as mock_warning:
        window._close_document()

    mock_warning.assert_not_called()
    assert window.tab_widget.count() == 0
    assert window.controller is None
    window.close()


def test_closing_a_dirty_document_prompts_and_cancel_keeps_it_open(
    qapp: QApplication, tmp_path: Path
) -> None:
    from core.ops.organize import RotatePagesOperation

    src = _make_pdf(tmp_path / "src.pdf", 1)
    window = MainWindow()
    _open_tab(window, src)
    window.controller.apply_operation(RotatePagesOperation(angle=90))
    window._refresh()

    with patch.object(QMessageBox, "warning", return_value=QMessageBox.StandardButton.Cancel):
        window._close_document()

    assert window.tab_widget.count() == 1
    assert window.controller.is_open
    # window.close() alone would hang here: the document is still
    # (correctly) dirty, so it'd trigger a second, unmocked closeEvent
    # -> a real modal QMessageBox.warning() blocking forever
    # headlessly. Clear the session directly first, then close() is
    # safe (nothing left to prompt about).
    _force_close(window)


def test_closing_a_dirty_document_discard_closes_it(qapp: QApplication, tmp_path: Path) -> None:
    from core.ops.organize import RotatePagesOperation

    src = _make_pdf(tmp_path / "src.pdf", 1)
    window = MainWindow()
    _open_tab(window, src)
    window.controller.apply_operation(RotatePagesOperation(angle=90))
    window._refresh()

    with patch.object(QMessageBox, "warning", return_value=QMessageBox.StandardButton.Discard):
        window._close_document()

    assert window.tab_widget.count() == 0
    window.close()


def test_window_close_event_is_ignored_when_user_cancels_unsaved_prompt(
    qapp: QApplication, tmp_path: Path
) -> None:
    from core.ops.organize import RotatePagesOperation

    src = _make_pdf(tmp_path / "src.pdf", 1)
    window = MainWindow()
    _open_tab(window, src)
    window.controller.apply_operation(RotatePagesOperation(angle=90))
    window._refresh()

    event = QCloseEvent()
    with patch.object(QMessageBox, "warning", return_value=QMessageBox.StandardButton.Cancel):
        window.closeEvent(event)

    assert not event.isAccepted()
    assert window.controller.is_open
    # See the comment in test_closing_a_dirty_document_prompts_and_cancel_keeps_it_open
    # - window.close() alone here would hang on a real, unmocked prompt.
    _force_close(window)


# --- recent files -----------------------------------------------------------


def test_opening_a_document_adds_it_to_recent_files(qapp: QApplication, tmp_path: Path) -> None:
    src = _make_pdf(tmp_path / "src.pdf", 1)
    window = MainWindow()

    window._open_document_path(src)

    assert window.recent_files.list() == [src]
    _force_close(window)


def test_recent_files_menu_lists_most_recent_first_and_reopens_on_click(
    qapp: QApplication, tmp_path: Path
) -> None:
    a = _make_pdf(tmp_path / "a.pdf", 1)
    b = _make_pdf(tmp_path / "b.pdf", 2)
    window = MainWindow()
    _open_tab(window, a)
    _open_tab(window, b)

    window._populate_recent_files_menu()
    actions = window.recent_files_menu.actions()
    # newest first, then a separator, then "Clear Recent Files"
    assert [a.text() for a in actions[:2]] == ["b.pdf", "a.pdf"]

    # Reopening through the menu goes through the same New Tab /
    # Replace Current Tab prompt File > Open does.
    with patch.object(TabPlacementDialog, "exec", _fake_placement(PLACEMENT_NEW_TAB)):
        actions[1].trigger()  # reopen a.pdf

    assert window.controller.doc.source_path == a
    assert window.tab_widget.count() == 3
    _force_close(window)


def test_recent_files_menu_shows_placeholder_when_empty(qapp: QApplication) -> None:
    window = MainWindow()

    window._populate_recent_files_menu()

    actions = window.recent_files_menu.actions()
    assert len(actions) == 1
    assert not actions[0].isEnabled()
    window.close()


def test_opening_a_stale_recent_file_shows_error_and_drops_it_from_the_list(
    qapp: QApplication, tmp_path: Path
) -> None:
    missing = tmp_path / "gone.pdf"
    window = MainWindow()
    window.recent_files.add(missing)

    with patch.object(QMessageBox, "critical") as mock_critical:
        window._open_recent_file(missing)

    mock_critical.assert_called_once()
    assert window.recent_files.list() == []
    window.close()


def test_clear_recent_files_empties_the_list(qapp: QApplication, tmp_path: Path) -> None:
    src = _make_pdf(tmp_path / "src.pdf", 1)
    window = MainWindow()
    window._open_document_path(src)
    assert window.recent_files.list() == [src]

    window.recent_files.clear()

    assert window.recent_files.list() == []
    _force_close(window)


def test_opening_a_recent_file_over_a_dirty_document_prompts_first(
    qapp: QApplication, tmp_path: Path
) -> None:
    from core.ops.organize import RotatePagesOperation

    a = _make_pdf(tmp_path / "a.pdf", 1)
    b = _make_pdf(tmp_path / "b.pdf", 1)
    window = MainWindow()
    _open_tab(window, a)
    window.controller.apply_operation(RotatePagesOperation(angle=90))
    window.recent_files.add(b)

    # Replace Current Tab still has to clear the dirty check for the
    # tab it's about to overwrite - cancelling that leaves it alone.
    with (
        patch.object(TabPlacementDialog, "exec", _fake_placement(PLACEMENT_REPLACE_CURRENT)),
        patch.object(QMessageBox, "warning", return_value=QMessageBox.StandardButton.Cancel),
    ):
        window._open_recent_file(b)

    assert window.controller.doc.source_path == a
    assert window.tab_widget.count() == 1
    _force_close(window)


# --- thumbnail context menu --------------------------------------------------
#
# QMenu.exec() is a native/compiled PySide6 method - unlike the
# project's own BaseToolDialog subclasses (a real Python class, so
# `patch.object(SomeDialog, "exec", fake)` genuinely overrides it),
# patching it the same way does NOT intercept the call: menu.exec(pos)
# still runs the real blocking modal popup, which hangs forever with
# no headless UI to click through (confirmed - hung pytest, had to
# kill -9 the stuck process). So these tests exercise the actual
# operation-applying logic directly (_apply_thumbnail_rotate /
# _apply_thumbnail_delete) rather than trying to drive it through a
# faked context-menu popup.


def test_thumbnail_context_menu_rotate_right_rotates_selected_pages_only(
    qapp: QApplication, tmp_path: Path
) -> None:
    src = _make_pdf(tmp_path / "src.pdf", 3)
    window = MainWindow()
    tab = _open_tab(window, src)

    window._apply_thumbnail_rotate(tab, [2], angle=90)

    with pikepdf.Pdf.open(window.controller.doc.working_path) as pdf:
        assert int(pdf.pages[0].get("/Rotate", 0)) == 0
        assert int(pdf.pages[1].get("/Rotate", 0)) == 90
        assert int(pdf.pages[2].get("/Rotate", 0)) == 0
    _force_close(window)


def test_thumbnail_context_menu_rotate_left_uses_negative_angle(
    qapp: QApplication, tmp_path: Path
) -> None:
    src = _make_pdf(tmp_path / "src.pdf", 1)
    window = MainWindow()
    tab = _open_tab(window, src)

    window._apply_thumbnail_rotate(tab, [1], angle=-90)

    with pikepdf.Pdf.open(window.controller.doc.working_path) as pdf:
        assert int(pdf.pages[0].get("/Rotate", 0)) == 270
    _force_close(window)


def test_thumbnail_context_menu_delete_removes_selected_pages(
    qapp: QApplication, tmp_path: Path
) -> None:
    src = _make_pdf(tmp_path / "src.pdf", 3)
    window = MainWindow()
    tab = _open_tab(window, src)

    window._apply_thumbnail_delete(tab, [1, 3])

    with pikepdf.Pdf.open(window.controller.doc.working_path) as pdf:
        assert len(pdf.pages) == 1
    _force_close(window)


def test_thumbnail_context_menu_does_nothing_without_a_selection(
    qapp: QApplication, tmp_path: Path
) -> None:
    src = _make_pdf(tmp_path / "src.pdf", 1)
    window = MainWindow()
    tab = _open_tab(window, src)
    assert tab.thumbnail_list.selectedItems() == []

    with patch.object(QMenu, "exec") as mock_exec:
        window._show_thumbnail_context_menu(tab, QPoint(0, 0))

    mock_exec.assert_not_called()
    _force_close(window)


# --- busy-cursor feedback ----------------------------------------------------


def test_busy_cursor_sets_wait_cursor_and_restores_it_after(qapp: QApplication) -> None:
    window = MainWindow()
    assert QApplication.overrideCursor() is None

    with window._busy_cursor():
        assert QApplication.overrideCursor() is not None
        assert QApplication.overrideCursor().shape() == Qt.CursorShape.WaitCursor

    assert QApplication.overrideCursor() is None
    window.close()


def test_busy_cursor_restores_cursor_even_on_exception(qapp: QApplication) -> None:
    window = MainWindow()

    with pytest.raises(RuntimeError), window._busy_cursor():
        raise RuntimeError("boom")

    assert QApplication.overrideCursor() is None
    window.close()


def test_running_a_tool_leaves_no_override_cursor_set_afterward(
    qapp: QApplication, tmp_path: Path
) -> None:
    src = _make_pdf(tmp_path / "src.pdf", 1)
    window = MainWindow()
    _open_tab(window, src)

    def fake_exec(self: RotateDialog) -> QDialog.DialogCode:
        self.angle.setCurrentText("180")
        return QDialog.DialogCode.Accepted

    with patch.object(RotateDialog, "exec", fake_exec):
        window._run_tool("rotate_pages", RotateDialog)

    assert QApplication.overrideCursor() is None
    _force_close(window)


# --- Phase 5: Workflow builder + run --------------------------------------


def test_workflows_menu_actions_exist_and_are_enabled_without_a_document(
    qapp: QApplication,
) -> None:
    # Building/running a workflow is document-independent by design -
    # unlike every tool in tool_actions, these two aren't gated on
    # is_open.
    window = MainWindow()
    assert window.build_workflow_action.isEnabled()
    assert window.run_workflow_action.isEnabled()
    window.close()


def test_build_workflow_saves_a_real_pipeline(qapp: QApplication) -> None:
    from core.registry.registry import Registry, discover_and_load
    from core.session.workflow_store import WorkflowStore

    window = MainWindow()
    registry = Registry()
    discover_and_load(registry)
    rotate_op = registry.get("rotate_pages").build_operation(angle=90, pages=[])

    def fake_exec(self: WorkflowBuilderDialog) -> QDialog.DialogCode:
        self.name_edit.setText("gui_test_workflow")
        self._operations.append(rotate_op)
        self.step_list.addItem(rotate_op.describe())
        return QDialog.DialogCode.Accepted

    with patch.object(WorkflowBuilderDialog, "exec", fake_exec):
        window._build_workflow()

    assert "gui_test_workflow" in WorkflowStore().list_workflows()
    window.close()


def test_workflow_builder_add_step_excludes_fill_form(qapp: QApplication) -> None:
    from core.registry.registry import Registry, discover_and_load

    registry = Registry()
    discover_and_load(registry)
    dialog = WorkflowBuilderDialog(registry)

    with patch(
        "gui.dialogs.workflow_builder_dialog.QInputDialog.getItem", return_value=("", False)
    ) as mock_get_item:
        dialog._add_step()

    offered_labels = mock_get_item.call_args[0][3]
    fill_form_display_name = registry.get("fill_form").display_name
    assert fill_form_display_name not in offered_labels
    dialog.close()


def test_workflow_builder_add_step_builds_and_lists_a_real_operation(qapp: QApplication) -> None:
    from core.registry.registry import Registry, discover_and_load

    registry = Registry()
    discover_and_load(registry)
    dialog = WorkflowBuilderDialog(registry)
    rotate_display_name = registry.get("rotate_pages").display_name

    def fake_rotate_exec(self: RotateDialog) -> QDialog.DialogCode:
        self.angle.setCurrentText("90")
        return QDialog.DialogCode.Accepted

    with (
        patch(
            "gui.dialogs.workflow_builder_dialog.QInputDialog.getItem",
            return_value=(rotate_display_name, True),
        ),
        patch.object(RotateDialog, "exec", fake_rotate_exec),
    ):
        dialog._add_step()

    assert dialog.step_list.count() == 1
    assert len(dialog._operations) == 1
    assert dialog._operations[0].describe() == "Rotated all pages by 90 degrees"
    dialog.close()


def test_workflow_builder_move_up_keeps_list_and_operations_in_sync(qapp: QApplication) -> None:
    from core.registry.registry import Registry, discover_and_load

    registry = Registry()
    discover_and_load(registry)
    dialog = WorkflowBuilderDialog(registry)

    op_a = registry.get("rotate_pages").build_operation(angle=90, pages=[])
    op_b = registry.get("flip").build_operation(direction="horizontal", pages=[])
    dialog._operations.extend([op_a, op_b])
    dialog.step_list.addItem(op_a.describe())
    dialog.step_list.addItem(op_b.describe())

    dialog.step_list.setCurrentRow(1)
    dialog._move(-1)

    assert dialog._operations == [op_b, op_a]
    assert [dialog.step_list.item(i).text() for i in range(2)] == [
        op_b.describe(),
        op_a.describe(),
    ]
    dialog.close()


def test_workflow_builder_remove_selected_keeps_list_and_operations_in_sync(
    qapp: QApplication,
) -> None:
    from core.registry.registry import Registry, discover_and_load

    registry = Registry()
    discover_and_load(registry)
    dialog = WorkflowBuilderDialog(registry)

    op_a = registry.get("rotate_pages").build_operation(angle=90, pages=[])
    op_b = registry.get("flip").build_operation(direction="horizontal", pages=[])
    dialog._operations.extend([op_a, op_b])
    dialog.step_list.addItem(op_a.describe())
    dialog.step_list.addItem(op_b.describe())

    dialog.step_list.setCurrentRow(0)
    dialog._remove_selected()

    assert dialog._operations == [op_b]
    assert dialog.step_list.count() == 1
    assert dialog.step_list.item(0).text() == op_b.describe()
    dialog.close()


def test_workflow_builder_add_step_shows_error_and_does_not_add_on_invalid_operation(
    qapp: QApplication,
) -> None:
    # Regression coverage: _add_step catches PDFEditorError from
    # build_operation (e.g. a dialog that accepted invalid input) and
    # shows an error dialog rather than appending a broken step or
    # crashing. Exercised via "compress", whose real dialog takes no
    # input but whose build_operation is stubbed here to fail, so the
    # error path is proven without depending on any one tool's own
    # validation quirks.
    from core.errors import OperationError
    from core.registry.registry import Registry, discover_and_load
    from gui.dialogs.compress_dialog import CompressDialog

    registry = Registry()
    discover_and_load(registry)
    dialog = WorkflowBuilderDialog(registry)
    compress_display_name = registry.get("compress").display_name

    def fake_compress_exec(self: CompressDialog) -> QDialog.DialogCode:
        return QDialog.DialogCode.Accepted

    with (
        patch(
            "gui.dialogs.workflow_builder_dialog.QInputDialog.getItem",
            return_value=(compress_display_name, True),
        ),
        patch.object(CompressDialog, "exec", fake_compress_exec),
        patch.object(
            type(registry.get("compress")),
            "build_operation",
            side_effect=OperationError("boom"),
        ),
        patch("gui.dialogs.workflow_builder_dialog.QMessageBox.critical") as mock_critical,
    ):
        dialog._add_step()

    mock_critical.assert_called_once()
    assert dialog.step_list.count() == 0
    assert dialog._operations == []
    dialog.close()


def test_workflow_builder_accept_rejects_empty_name(qapp: QApplication) -> None:
    from core.registry.registry import Registry, discover_and_load

    registry = Registry()
    discover_and_load(registry)
    dialog = WorkflowBuilderDialog(registry)
    op = registry.get("rotate_pages").build_operation(angle=90, pages=[])
    dialog._operations.append(op)
    dialog.step_list.addItem(op.describe())

    with patch("gui.dialogs.workflow_builder_dialog.QMessageBox.warning") as mock_warning:
        dialog.accept()
    mock_warning.assert_called_once()
    assert dialog.result() != QDialog.DialogCode.Accepted
    dialog.close()


def test_workflow_builder_accept_rejects_zero_steps(qapp: QApplication) -> None:
    from core.registry.registry import Registry, discover_and_load

    registry = Registry()
    discover_and_load(registry)
    dialog = WorkflowBuilderDialog(registry)
    dialog.name_edit.setText("has_a_name")

    with patch("gui.dialogs.workflow_builder_dialog.QMessageBox.warning") as mock_warning:
        dialog.accept()
    mock_warning.assert_called_once()
    dialog.close()


def test_run_workflow_shows_info_when_no_workflows_saved(qapp: QApplication) -> None:
    window = MainWindow()
    with patch("gui.main_window.QMessageBox.information") as mock_info:
        window._run_workflow()
    mock_info.assert_called_once()
    window.close()


def test_run_workflow_applies_saved_pipeline_without_touching_open_document(
    qapp: QApplication, tmp_path: Path
) -> None:
    from core.model.pipeline import Pipeline
    from core.registry.registry import Registry, discover_and_load
    from core.session.workflow_store import WorkflowStore

    registry = Registry()
    discover_and_load(registry)
    rotate = registry.get("rotate_pages").build_operation(angle=90, pages=[])
    WorkflowStore().save(Pipeline(name="gui_run_test", operations=[rotate]))

    src = _make_pdf(tmp_path / "src.pdf", 1)
    out = tmp_path / "out.pdf"

    window = MainWindow()
    # a currently-open, unrelated document - Run Workflow's batch
    # semantics must leave it completely untouched.
    other_open = _make_pdf(tmp_path / "other.pdf", 1)
    _open_tab(window, other_open)
    ops_before = len(window.controller.doc.operation_log)

    def fake_exec(self: RunWorkflowDialog) -> QDialog.DialogCode:
        self.workflow_combo.setCurrentText("gui_run_test")
        self._input_path = src
        self._output_path = out
        return QDialog.DialogCode.Accepted

    with patch.object(RunWorkflowDialog, "exec", fake_exec):
        window._run_workflow()

    assert out.exists()
    with pikepdf.Pdf.open(out) as pdf:
        assert int(pdf.pages[0].get("/Rotate", 0)) == 90
    assert len(window.controller.doc.operation_log) == ops_before
    _force_close(window)


def test_view_menu_zoom_in_out_and_reset_resize_the_icon_and_rerender(
    qapp: QApplication, tmp_path: Path
) -> None:
    from gui.main_window import _THUMBNAIL_SIZE

    src = _make_pdf(tmp_path / "src.pdf", 1)
    window = MainWindow()
    _open_tab(window, src)

    assert window.thumbnail_size == _THUMBNAIL_SIZE
    assert window.thumbnail_list.iconSize() == _THUMBNAIL_SIZE

    window.zoom_in_action.trigger()

    assert window.thumbnail_size.width() > _THUMBNAIL_SIZE.width()
    assert window.thumbnail_list.iconSize() == window.thumbnail_size
    # Not just an internal size variable - the actual rendered icon
    # pixmap must be re-rendered at the new size, not stretched.
    item = window.thumbnail_list.item(0)
    assert item.icon().actualSize(window.thumbnail_size) == window.thumbnail_size

    window.zoom_out_action.trigger()
    window.zoom_out_action.trigger()

    assert window.thumbnail_size.width() < _THUMBNAIL_SIZE.width()
    assert window.thumbnail_list.iconSize() == window.thumbnail_size
    item = window.thumbnail_list.item(0)
    assert item.icon().actualSize(window.thumbnail_size) == window.thumbnail_size

    window.reset_zoom_action.trigger()

    assert window.thumbnail_size == _THUMBNAIL_SIZE
    assert window.thumbnail_list.iconSize() == _THUMBNAIL_SIZE

    _force_close(window)


def test_view_menu_zoom_is_clamped_to_min_and_max(qapp: QApplication) -> None:
    from gui.main_window import _THUMBNAIL_ZOOM_MAX_WIDTH, _THUMBNAIL_ZOOM_MIN_WIDTH

    window = MainWindow()

    for _ in range(50):
        window.zoom_in_action.trigger()
    assert window.thumbnail_size.width() == _THUMBNAIL_ZOOM_MAX_WIDTH

    for _ in range(50):
        window.zoom_out_action.trigger()
    assert window.thumbnail_size.width() == _THUMBNAIL_ZOOM_MIN_WIDTH

    window.close()


def test_view_menu_zoom_in_keyboard_shortcut_actually_fires(
    qapp: QApplication, tmp_path: Path
) -> None:
    """Regression test for a real user report: "Ctrl++ isn't working."

    QKeySequence.StandardKey.ZoomIn resolves to the literal "Ctrl++",
    but '+' is Shift+'=' on a US keyboard layout (and varies further on
    non-US ones) - a user pressing the unshifted "Ctrl+=" saw nothing
    happen. This drives real QTest key events (not `.trigger()`, which
    would pass even if no shortcut were bound at all) through the
    actual QAction shortcut-matching machinery, covering both the
    literal Ctrl++ (still bound, must keep working) and the added
    unshifted Ctrl+= alternate.
    """
    from gui.main_window import _THUMBNAIL_ZOOM_STEP

    src = _make_pdf(tmp_path / "src.pdf", 1)
    window = MainWindow()
    window.show()
    _open_tab(window, src)
    # QTest key events are only routed to shortcuts when the window is
    # the active one - confirmed by hand while investigating this bug
    # (a shown-but-not-activated offscreen window silently drops every
    # QTest.keyClick-driven shortcut, independent of this fix).
    window.activateWindow()
    QApplication.processEvents()

    start_width = window.thumbnail_size.width()

    QTest.keyClick(window, Qt.Key.Key_Equal, Qt.KeyboardModifier.ControlModifier)
    assert window.thumbnail_size.width() == start_width + _THUMBNAIL_ZOOM_STEP

    QTest.keyClick(window, Qt.Key.Key_Plus, Qt.KeyboardModifier.ControlModifier)
    assert window.thumbnail_size.width() == start_width + 2 * _THUMBNAIL_ZOOM_STEP

    _force_close(window)


def test_view_menu_toggle_toolbar_visibility(qapp: QApplication) -> None:
    window = MainWindow()
    window.show()
    assert window.toolbar.isVisible()
    assert window.toggle_toolbar_action.isChecked()

    window.toggle_toolbar_action.trigger()
    assert not window.toolbar.isVisible()
    assert not window.toggle_toolbar_action.isChecked()

    window.toggle_toolbar_action.trigger()
    assert window.toolbar.isVisible()
    window.close()


def test_view_menu_toggle_status_bar_visibility(qapp: QApplication) -> None:
    window = MainWindow()
    window.show()
    assert window.statusBar().isVisible()
    assert window.toggle_statusbar_action.isChecked()

    window.toggle_statusbar_action.trigger()
    assert not window.statusBar().isVisible()
    assert not window.toggle_statusbar_action.isChecked()

    window.toggle_statusbar_action.trigger()
    assert window.statusBar().isVisible()
    window.close()


def test_view_menu_full_screen_toggle_reflects_window_state(qapp: QApplication) -> None:
    # Under QT_QPA_PLATFORM=offscreen, showFullScreen()/showNormal()
    # were confirmed by hand to actually flip Qt.WindowState.WindowFullScreen
    # (not a headless no-op) - this test exercises the real toggle, not
    # a stand-in for one.
    window = MainWindow()
    window.show()
    assert not window.isFullScreen()

    window.full_screen_action.trigger()
    assert window.isFullScreen()
    assert window.full_screen_action.isChecked()

    window.full_screen_action.trigger()
    assert not window.isFullScreen()
    assert not window.full_screen_action.isChecked()
    window.close()


def test_run_workflow_records_every_step_in_the_audit_log(
    qapp: QApplication, tmp_path: Path
) -> None:
    # Regression: unlike every other path that applies an Operation
    # (AppController.apply_operation, and the CLI's own run-workflow,
    # see test_run_workflow_records_every_step_in_the_audit_log in
    # tests/integration/test_cli.py), MainWindow._run_workflow used to
    # not record anything to the audit trail at all - a GUI-driven
    # workflow run was invisible to it despite genuinely modifying a
    # document.
    from core.model.pipeline import Pipeline
    from core.registry.registry import Registry, discover_and_load
    from core.session.audit_log import AuditLog
    from core.session.workflow_store import WorkflowStore

    registry = Registry()
    discover_and_load(registry)
    rotate = registry.get("rotate_pages").build_operation(angle=90, pages=[])
    watermark = registry.get("watermark").build_operation(
        text="DRAFT", opacity=0.3, font_size=40
    )
    WorkflowStore().save(Pipeline(name="gui_audit_test", operations=[rotate, watermark]))

    src = _make_pdf(tmp_path / "src.pdf", 1)
    out = tmp_path / "out.pdf"

    window = MainWindow()

    def fake_exec(self: RunWorkflowDialog) -> QDialog.DialogCode:
        self.workflow_combo.setCurrentText("gui_audit_test")
        self._input_path = src
        self._output_path = out
        return QDialog.DialogCode.Accepted

    with patch.object(RunWorkflowDialog, "exec", fake_exec):
        window._run_workflow()

    entries = AuditLog().read_all()
    assert [e["operation"]["type"] for e in entries] == ["rotate_pages", "watermark"]
    assert all(e["document"] == str(out) for e in entries)
    window.close()


# --- multi-document tabs -----------------------------------------------------
#
# The point of these is per-tab *isolation* being real, not just
# rendered: an operation applied in one tab must leave the other tab's
# DocumentSession, undo/redo stack, working file and session temp dir
# genuinely untouched, checked against the actual PDFs and directories
# rather than against the UI's own idea of what happened.


def test_two_tabs_have_genuinely_independent_undo_stacks(
    qapp: QApplication, tmp_path: Path
) -> None:
    from core.ops.organize import RotatePagesOperation

    a = _make_pdf(tmp_path / "a.pdf", 2)
    b = _make_pdf(tmp_path / "b.pdf", 3)
    window = MainWindow()
    tab_a = _open_tab(window, a)
    tab_b = _open_tab(window, b)

    assert tab_a is not tab_b
    assert tab_a.controller is not tab_b.controller
    # Separate private session dirs, not one shared scratch space.
    assert tab_a.controller.doc.working_path.parent != tab_b.controller.doc.working_path.parent

    window.tab_widget.setCurrentWidget(tab_a)
    window.controller.apply_operation(RotatePagesOperation(angle=90))
    window._refresh()

    # Tab A changed...
    with pikepdf.Pdf.open(tab_a.controller.doc.working_path) as pdf:
        assert int(pdf.pages[0].get("/Rotate", 0)) == 90
    assert len(tab_a.controller.doc.operation_log) == 1
    # ...and tab B did not, at every level: log, redo stack, and the
    # actual bytes of its own working file.
    assert tab_b.controller.doc.operation_log == []
    assert tab_b.controller.doc.redo_stack == []
    assert not tab_b.controller.is_dirty
    with pikepdf.Pdf.open(tab_b.controller.doc.working_path) as pdf:
        assert len(pdf.pages) == 3
        assert int(pdf.pages[0].get("/Rotate", 0)) == 0

    window._undo()

    assert tab_a.controller.doc.operation_log == []
    assert len(tab_a.controller.doc.redo_stack) == 1
    with pikepdf.Pdf.open(tab_a.controller.doc.working_path) as pdf:
        assert int(pdf.pages[0].get("/Rotate", 0)) == 0
    # The undo in A must not have reached into B's stack either.
    assert tab_b.controller.doc.operation_log == []
    assert tab_b.controller.doc.redo_stack == []

    # Undo/redo actions reflect whichever tab is active.
    window.tab_widget.setCurrentWidget(tab_b)
    assert not window.undo_action.isEnabled()
    assert not window.redo_action.isEnabled()
    window.tab_widget.setCurrentWidget(tab_a)
    assert window.redo_action.isEnabled()

    _force_close(window)


def test_each_tab_tracks_its_own_dirty_state_and_tab_label(
    qapp: QApplication, tmp_path: Path
) -> None:
    from core.ops.organize import RotatePagesOperation

    a = _make_pdf(tmp_path / "a.pdf", 1)
    b = _make_pdf(tmp_path / "b.pdf", 1)
    window = MainWindow()
    tab_a = _open_tab(window, a)
    tab_b = _open_tab(window, b)

    assert window.tab_widget.tabText(0) == "a.pdf"
    assert window.tab_widget.tabText(1) == "b.pdf"

    window.tab_widget.setCurrentWidget(tab_a)
    window.controller.apply_operation(RotatePagesOperation(angle=90))
    window._refresh()

    assert tab_a.controller.is_dirty
    assert not tab_b.controller.is_dirty
    assert window.tab_widget.tabText(0).startswith("•")
    assert "a.pdf" in window.tab_widget.tabText(0)
    assert window.tab_widget.tabText(1) == "b.pdf"

    out = tmp_path / "saved.pdf"
    with patch("gui.main_window.QFileDialog.getSaveFileName", return_value=(str(out), "")):
        assert window._save_as(tab_a)

    assert not tab_a.controller.is_dirty
    assert window.tab_widget.tabText(0) == "a.pdf"
    _force_close(window)


def test_closing_a_tab_wipes_only_that_tabs_session(qapp: QApplication, tmp_path: Path) -> None:
    from core.ops.organize import RotatePagesOperation

    a = _make_pdf(tmp_path / "a.pdf", 1)
    b = _make_pdf(tmp_path / "b.pdf", 2)
    window = MainWindow()
    tab_a = _open_tab(window, a)
    tab_b = _open_tab(window, b)
    session_a = tab_a.controller.doc.working_path.parent
    session_b = tab_b.controller.doc.working_path.parent
    assert session_a.exists()
    assert session_b.exists()

    window._close_tab(window.tab_widget.indexOf(tab_a))

    assert not session_a.exists()  # securely wiped now, not at app exit
    assert session_b.exists()
    assert window.tab_widget.count() == 1
    assert window.current_tab is tab_b

    # The surviving tab still works normally afterwards - the wipe
    # didn't take anything it needed with it.
    window.controller.apply_operation(RotatePagesOperation(angle=180))
    window._refresh()
    with pikepdf.Pdf.open(tab_b.controller.doc.working_path) as pdf:
        assert len(pdf.pages) == 2
        assert int(pdf.pages[0].get("/Rotate", 0)) == 180
    assert window.thumbnail_list.count() == 2
    _force_close(window)


def test_close_other_tabs_dirty_checks_each_closed_tab(
    qapp: QApplication, tmp_path: Path
) -> None:
    from core.ops.organize import RotatePagesOperation

    a = _make_pdf(tmp_path / "a.pdf", 1)
    b = _make_pdf(tmp_path / "b.pdf", 1)
    c = _make_pdf(tmp_path / "c.pdf", 1)
    window = MainWindow()
    tab_a = _open_tab(window, a)
    tab_b = _open_tab(window, b)
    tab_c = _open_tab(window, c)

    # A mix of clean and dirty: only the dirty one should be asked about.
    window.tab_widget.setCurrentWidget(tab_a)
    window.controller.apply_operation(RotatePagesOperation(angle=90))
    session_a = tab_a.controller.doc.working_path.parent
    session_c = tab_c.controller.doc.working_path.parent

    with patch.object(
        QMessageBox, "warning", return_value=QMessageBox.StandardButton.Discard
    ) as mock_warning:
        assert window._close_other_tabs(window.tab_widget.indexOf(tab_b))

    mock_warning.assert_called_once()  # tab_a only; tab_c was clean
    assert window.tab_widget.count() == 1
    assert window.current_tab is tab_b
    assert not session_a.exists()
    assert not session_c.exists()
    _force_close(window)


def test_close_all_tabs_cancelled_on_a_dirty_tab_keeps_it_open(
    qapp: QApplication, tmp_path: Path
) -> None:
    from core.ops.organize import RotatePagesOperation

    a = _make_pdf(tmp_path / "a.pdf", 1)
    b = _make_pdf(tmp_path / "b.pdf", 1)
    window = MainWindow()
    _open_tab(window, a)
    tab_b = _open_tab(window, b)
    window.controller.apply_operation(RotatePagesOperation(angle=90))

    with patch.object(QMessageBox, "warning", return_value=QMessageBox.StandardButton.Cancel):
        assert not window._close_all_tabs()

    # The clean tab ahead of it did close; the cancelled one stayed.
    assert window.tab_widget.count() == 1
    assert window.current_tab is tab_b
    assert tab_b.controller.is_dirty
    _force_close(window)


def test_window_close_checks_every_tab_not_only_the_active_one(
    qapp: QApplication, tmp_path: Path
) -> None:
    from core.ops.organize import RotatePagesOperation

    a = _make_pdf(tmp_path / "a.pdf", 1)
    b = _make_pdf(tmp_path / "b.pdf", 1)
    window = MainWindow()
    tab_a = _open_tab(window, a)
    tab_b = _open_tab(window, b)

    # Dirty the *background* tab, then leave a clean tab active: the
    # old single-document closeEvent would have seen nothing to lose.
    window.tab_widget.setCurrentWidget(tab_a)
    window.controller.apply_operation(RotatePagesOperation(angle=90))
    window.tab_widget.setCurrentWidget(tab_b)
    assert not window.controller.is_dirty

    event = QCloseEvent()
    with patch.object(
        QMessageBox, "warning", return_value=QMessageBox.StandardButton.Cancel
    ) as mock_warning:
        window.closeEvent(event)

    mock_warning.assert_called_once()
    assert not event.isAccepted()
    assert tab_a in window.tabs()
    _force_close(window)


def test_tabs_can_be_reordered_and_keep_their_own_documents(
    qapp: QApplication, tmp_path: Path
) -> None:
    # tabBar().moveTab(...) is the same call Qt makes at the end of a
    # real drag-to-reorder gesture - real drags aren't reliably
    # simulatable under QT_QPA_PLATFORM=offscreen (see
    # test_dragging_a_thumbnail_reorders_the_document).
    a = _make_pdf(tmp_path / "a.pdf", 1)
    b = _make_pdf(tmp_path / "b.pdf", 2)
    window = MainWindow()
    tab_a = _open_tab(window, a)
    tab_b = _open_tab(window, b)
    assert window.tab_widget.isMovable()
    assert window.tabs() == [tab_a, tab_b]

    window.tab_widget.tabBar().moveTab(0, 1)

    assert window.tabs() == [tab_b, tab_a]
    assert window.tab_widget.tabText(0) == "b.pdf"
    assert window.tab_widget.tabText(1) == "a.pdf"
    # Reordering is presentation only - each tab still owns its own
    # document and working file.
    assert tab_a.controller.doc.source_path == a
    assert tab_b.controller.doc.source_path == b
    with pikepdf.Pdf.open(tab_b.controller.doc.working_path) as pdf:
        assert len(pdf.pages) == 2
    _force_close(window)


def test_ctrl_tab_cycles_through_tabs_in_both_directions(
    qapp: QApplication, tmp_path: Path
) -> None:
    paths = [
        _make_pdf(tmp_path / "a.pdf", 1),
        _make_pdf(tmp_path / "b.pdf", 1),
        _make_pdf(tmp_path / "c.pdf", 1),
    ]
    window = MainWindow()
    for path in paths:
        _open_tab(window, path)

    assert window.next_tab_action.shortcut().toString() == "Ctrl+Tab"
    assert window.previous_tab_action.shortcut().toString() == "Ctrl+Shift+Tab"

    window.tab_widget.setCurrentIndex(0)
    window.next_tab_action.trigger()
    assert window.tab_widget.currentIndex() == 1
    window.next_tab_action.trigger()
    window.next_tab_action.trigger()
    assert window.tab_widget.currentIndex() == 0  # wraps around

    window.previous_tab_action.trigger()
    assert window.tab_widget.currentIndex() == 2  # wraps the other way
    _force_close(window)


def test_ctrl_w_closes_only_the_current_tab(qapp: QApplication, tmp_path: Path) -> None:
    a = _make_pdf(tmp_path / "a.pdf", 1)
    b = _make_pdf(tmp_path / "b.pdf", 1)
    window = MainWindow()
    tab_a = _open_tab(window, a)
    tab_b = _open_tab(window, b)

    assert window.close_action.shortcut().toString() == "Ctrl+W"
    window.tab_widget.setCurrentWidget(tab_a)
    window.close_action.trigger()

    assert window.tabs() == [tab_b]
    _force_close(window)


def test_opening_with_no_tabs_open_skips_the_placement_prompt(
    qapp: QApplication, tmp_path: Path
) -> None:
    src = _make_pdf(tmp_path / "src.pdf", 1)
    window = MainWindow()

    with patch.object(TabPlacementDialog, "exec") as mock_exec:
        window._open_document_path(src)

    mock_exec.assert_not_called()  # nothing to replace, nothing ambiguous
    assert window.tab_widget.count() == 1
    assert window.controller.doc.source_path == src
    _force_close(window)


def test_opening_a_second_document_in_a_new_tab_keeps_the_first(
    qapp: QApplication, tmp_path: Path
) -> None:
    a = _make_pdf(tmp_path / "a.pdf", 1)
    b = _make_pdf(tmp_path / "b.pdf", 2)
    window = MainWindow()
    tab_a = _open_tab(window, a)
    session_a = tab_a.controller.doc.working_path.parent

    with patch.object(TabPlacementDialog, "exec", _fake_placement(PLACEMENT_NEW_TAB)):
        window._open_document_path(b)

    assert window.tab_widget.count() == 2
    assert tab_a.controller.doc.source_path == a
    assert session_a.exists()
    assert window.controller.doc.source_path == b
    _force_close(window)


def test_replace_current_tab_replaces_only_that_tab(qapp: QApplication, tmp_path: Path) -> None:
    a = _make_pdf(tmp_path / "a.pdf", 1)
    b = _make_pdf(tmp_path / "b.pdf", 1)
    c = _make_pdf(tmp_path / "c.pdf", 3)
    window = MainWindow()
    tab_a = _open_tab(window, a)
    tab_b = _open_tab(window, b)
    old_session_b = tab_b.controller.doc.working_path.parent

    with patch.object(TabPlacementDialog, "exec", _fake_placement(PLACEMENT_REPLACE_CURRENT)):
        window._open_document_path(c)

    assert window.tab_widget.count() == 2
    assert window.current_tab is tab_b
    assert tab_b.controller.doc.source_path == c
    assert not old_session_b.exists()  # the replaced document's scratch space is wiped
    assert tab_a.controller.doc.source_path == a  # untouched
    assert window.tab_widget.tabText(window.tab_widget.indexOf(tab_b)) == "c.pdf"
    _force_close(window)


def test_cancelling_the_placement_prompt_opens_nothing(qapp: QApplication, tmp_path: Path) -> None:
    a = _make_pdf(tmp_path / "a.pdf", 1)
    b = _make_pdf(tmp_path / "b.pdf", 1)
    window = MainWindow()
    _open_tab(window, a)

    def fake_cancel(self: TabPlacementDialog) -> QDialog.DialogCode:
        return QDialog.DialogCode.Rejected

    with patch.object(TabPlacementDialog, "exec", fake_cancel):
        window._open_document_path(b)

    assert window.tab_widget.count() == 1
    assert window.controller.doc.source_path == a
    assert window.recent_files.list() == [a]
    _force_close(window)


def test_a_failed_open_in_a_new_tab_does_not_strand_an_empty_tab(
    qapp: QApplication, tmp_path: Path
) -> None:
    a = _make_pdf(tmp_path / "a.pdf", 1)
    window = MainWindow()
    _open_tab(window, a)

    with (
        patch.object(TabPlacementDialog, "exec", _fake_placement(PLACEMENT_NEW_TAB)),
        patch.object(QMessageBox, "critical") as mock_critical,
    ):
        window._open_document_path(tmp_path / "does-not-exist.pdf")

    mock_critical.assert_called_once()
    assert window.tab_widget.count() == 1
    assert window.controller.doc.source_path == a
    _force_close(window)


# --- black-empty-tab regression --------------------------------------------
#
# Real bug, found and fixed: creating a new tab made it *current*
# synchronously (QTabWidget.setCurrentIndex fires currentChanged
# immediately), which rendered it - via the normal _refresh() path -
# before the caller had actually opened/built a document in it. That
# produced a real, capturable frame: an empty "Untitled" tab with a
# plain dark thumbnail grid (zero items) and "0 page(s)" in the status
# bar, confirmed by grab()ing the real window under
# QT_QPA_PLATFORM=offscreen. The fix (see MainWindow._add_tab's
# docstring) defers activating a new tab until it actually has a
# document. These tests cover the root cause directly (activation is
# deferred) and the user-visible symptom (every tab's thumbnails are
# real, correctly-colored pixels, not a blank/stale grid), across all
# three tab-creation paths plus repeated switching.


def test_opening_a_new_tab_does_not_activate_it_until_its_document_is_loaded(
    qapp: QApplication, tmp_path: Path
) -> None:
    """Root-cause regression: while a second-or-later tab's document
    is being opened, the window must still be showing the *previous*
    tab, not the new, as-yet-empty one."""
    a = _make_pdf(tmp_path / "a.pdf", 1)
    b = _make_pdf(tmp_path / "b.pdf", 1)
    window = MainWindow()
    tab_a = _open_tab(window, a)

    observed: list[bool] = []
    original_open = tab_a.controller.__class__.open_document

    def spying_open_document(self: Any, path: Path) -> None:
        # Captured mid-call, before the real open_document has done
        # anything: the window must still be showing tab_a, not a
        # freshly-added-but-undocumented new tab.
        observed.append(window.current_tab is tab_a)
        original_open(self, path)

    with patch.object(type(tab_a.controller), "open_document", spying_open_document):
        window._open_document_path(b, PLACEMENT_NEW_TAB)

    assert observed == [True]
    # ...and once the document is actually loaded, the new tab (not A)
    # is the one that's current.
    assert window.current_tab is not tab_a
    assert window.current_tab is not None
    assert window.current_tab.document_name() == "b.pdf"
    _force_close(window)


def test_the_first_tab_ever_does_not_activate_until_its_document_is_loaded(
    qapp: QApplication, tmp_path: Path
) -> None:
    """Same root-cause check for the very first tab (zero tabs open):
    Qt auto-selects the first tab added to an empty QTabWidget even
    without an explicit setCurrentIndex call (confirmed directly), so
    this path needs its own coverage - the fix has to suppress that
    signal too, not just skip the manual setCurrentIndex(). While the
    document is loading, the window must still show the empty-state
    welcome screen, not a black tab."""
    a = _make_pdf(tmp_path / "a.pdf", 1)
    window = MainWindow()
    assert window.tab_widget.count() == 0

    observed: list[bool] = []

    def spying_open_document(self: Any, path: Path) -> None:
        observed.append(window.stack.currentWidget() is window.empty_state)

    with patch.object(AppController, "open_document", spying_open_document):
        window._open_document_path(a, PLACEMENT_NEW_TAB)

    assert observed == [True]
    _force_close(window)


def test_opening_a_second_tab_shows_its_own_real_page_content_by_pixel(
    qapp: QApplication, tmp_path: Path
) -> None:
    a = _make_colored_pdf(tmp_path / "a.pdf", 1, (1, 0, 0))
    b = _make_colored_pdf(tmp_path / "b.pdf", 1, (0, 1, 0))
    window = MainWindow()
    _open_tab(window, a)
    tab_b = _open_tab(window, b)

    assert tab_b.thumbnail_list.count() == 1
    assert _thumbnail_center_pixel(tab_b, window)[:3] == (0, 255, 0)
    _force_close(window)


def test_switching_back_to_a_backgrounded_tab_still_shows_its_real_content(
    qapp: QApplication, tmp_path: Path
) -> None:
    """Repeated switching (per the bug report's correction - the
    black tab might not be the newly-created one, it might be one
    reactivated after going to the background) - every switch, both
    directions, several times, must keep showing the real page color,
    never a blank/black grid."""
    a = _make_colored_pdf(tmp_path / "a.pdf", 1, (1, 0, 0))
    b = _make_colored_pdf(tmp_path / "b.pdf", 1, (0, 1, 0))
    window = MainWindow()
    tab_a = _open_tab(window, a)
    tab_b = _open_tab(window, b)

    for _ in range(5):
        window.tab_widget.setCurrentWidget(tab_a)
        assert tab_a.thumbnail_list.count() == 1
        assert _thumbnail_center_pixel(tab_a, window)[:3] == (255, 0, 0)
        window.tab_widget.setCurrentWidget(tab_b)
        assert tab_b.thumbnail_list.count() == 1
        assert _thumbnail_center_pixel(tab_b, window)[:3] == (0, 255, 0)
    _force_close(window)


def test_a_third_tab_and_switching_out_of_order_keeps_each_tabs_content_correct(
    qapp: QApplication, tmp_path: Path
) -> None:
    a = _make_colored_pdf(tmp_path / "a.pdf", 1, (1, 0, 0))
    b = _make_colored_pdf(tmp_path / "b.pdf", 1, (0, 1, 0))
    c = _make_colored_pdf(tmp_path / "c.pdf", 1, (0, 0, 1))
    window = MainWindow()
    tab_a = _open_tab(window, a)
    tab_b = _open_tab(window, b)
    tab_c = _open_tab(window, c)

    for tab, expected in [(tab_a, (255, 0, 0)), (tab_c, (0, 0, 255)), (tab_b, (0, 255, 0))]:
        window.tab_widget.setCurrentWidget(tab)
        assert _thumbnail_center_pixel(tab, window)[:3] == expected
    _force_close(window)


def test_replacing_the_current_tab_shows_the_new_documents_real_content(
    qapp: QApplication, tmp_path: Path
) -> None:
    a = _make_colored_pdf(tmp_path / "a.pdf", 1, (1, 0, 0))
    b = _make_colored_pdf(tmp_path / "b.pdf", 1, (0, 0, 1))
    window = MainWindow()
    _open_tab(window, a)

    window._open_document_path(b, PLACEMENT_REPLACE_CURRENT)

    assert window.tab_widget.count() == 1
    tab = window.current_tab
    assert tab is not None
    assert _thumbnail_center_pixel(tab, window)[:3] == (0, 0, 255)
    _force_close(window)


def test_a_merge_that_fails_to_build_does_not_strand_an_empty_tab(
    qapp: QApplication, tmp_path: Path
) -> None:
    """Merge with no document open gets its tab created only once the
    dialog is accepted (existing behavior for a *cancelled* dialog) -
    this covers the related gap found while fixing the black-tab bug:
    an *accepted* dialog whose build then fails (a missing input file)
    must not leave an empty, permanently-blank tab behind either."""
    window = MainWindow()
    assert window.tab_widget.count() == 0

    def fake_exec(self: MergeDialog) -> QDialog.DialogCode:
        self.file_list.addItems([str(tmp_path / "does-not-exist.pdf")])
        return QDialog.DialogCode.Accepted

    with (
        patch.object(MergeDialog, "exec", fake_exec),
        patch.object(QMessageBox, "critical") as mock_critical,
    ):
        window._run_tool("merge", MergeDialog)

    mock_critical.assert_called_once()
    assert window.tab_widget.count() == 0
    assert window.current_tab is None
    _force_close(window)


def test_a_tool_applies_only_to_the_active_tab(qapp: QApplication, tmp_path: Path) -> None:
    a = _make_pdf(tmp_path / "a.pdf", 1)
    b = _make_pdf(tmp_path / "b.pdf", 1)
    window = MainWindow()
    tab_a = _open_tab(window, a)
    tab_b = _open_tab(window, b)

    def fake_exec(self: RotateDialog) -> QDialog.DialogCode:
        self.angle.setCurrentText("180")
        return QDialog.DialogCode.Accepted

    window.tab_widget.setCurrentWidget(tab_b)
    with patch.object(RotateDialog, "exec", fake_exec):
        window._run_tool("rotate_pages", RotateDialog)

    with pikepdf.Pdf.open(tab_b.controller.doc.working_path) as pdf:
        assert int(pdf.pages[0].get("/Rotate", 0)) == 180
    with pikepdf.Pdf.open(tab_a.controller.doc.working_path) as pdf:
        assert int(pdf.pages[0].get("/Rotate", 0)) == 0
    assert tab_a.controller.doc.operation_log == []
    _force_close(window)


def test_thumbnail_context_menu_acts_on_the_tab_it_came_from(
    qapp: QApplication, tmp_path: Path
) -> None:
    a = _make_pdf(tmp_path / "a.pdf", 2)
    b = _make_pdf(tmp_path / "b.pdf", 2)
    window = MainWindow()
    tab_a = _open_tab(window, a)
    tab_b = _open_tab(window, b)

    window._apply_thumbnail_delete(tab_a, [1])

    with pikepdf.Pdf.open(tab_a.controller.doc.working_path) as pdf:
        assert len(pdf.pages) == 1
    with pikepdf.Pdf.open(tab_b.controller.doc.working_path) as pdf:
        assert len(pdf.pages) == 2
    _force_close(window)


def test_a_new_tab_uses_the_current_window_level_zoom(qapp: QApplication, tmp_path: Path) -> None:
    # Decision: zoom is window-level, not per-tab.
    a = _make_pdf(tmp_path / "a.pdf", 1)
    b = _make_pdf(tmp_path / "b.pdf", 1)
    window = MainWindow()
    tab_a = _open_tab(window, a)
    window.zoom_in_action.trigger()
    zoomed = QSize(window.thumbnail_size)

    tab_b = _open_tab(window, b)

    assert tab_b.thumbnail_list.iconSize() == zoomed
    assert tab_b.thumbnail_list.item(0).icon().actualSize(zoomed) == zoomed
    # ...and switching back re-renders the first tab at the same size,
    # rather than leaving it at the size it was opened with.
    window.tab_widget.setCurrentWidget(tab_a)
    assert tab_a.thumbnail_list.iconSize() == zoomed
    assert tab_a.thumbnail_list.item(0).icon().actualSize(zoomed) == zoomed
    _force_close(window)


def test_run_workflow_touches_no_tab_and_opens_none(qapp: QApplication, tmp_path: Path) -> None:
    from core.model.pipeline import Pipeline
    from core.registry.registry import Registry, discover_and_load
    from core.session.workflow_store import WorkflowStore

    registry = Registry()
    discover_and_load(registry)
    rotate = registry.get("rotate_pages").build_operation(angle=90, pages=[])
    WorkflowStore().save(Pipeline(name="gui_multitab_run", operations=[rotate]))

    src = _make_pdf(tmp_path / "src.pdf", 1)
    out = tmp_path / "out.pdf"

    window = MainWindow()
    tab_a = _open_tab(window, _make_pdf(tmp_path / "a.pdf", 1))
    tab_b = _open_tab(window, _make_pdf(tmp_path / "b.pdf", 1))

    def fake_exec(self: RunWorkflowDialog) -> QDialog.DialogCode:
        self.workflow_combo.setCurrentText("gui_multitab_run")
        self._input_path = src
        self._output_path = out
        return QDialog.DialogCode.Accepted

    with patch.object(RunWorkflowDialog, "exec", fake_exec):
        window._run_workflow()

    assert out.exists()
    with pikepdf.Pdf.open(out) as pdf:
        assert int(pdf.pages[0].get("/Rotate", 0)) == 90
    # No tab opened, and no tab's document, undo stack or dirty flag
    # touched - Run Workflow is batch replay, not live editing.
    assert window.tab_widget.count() == 2
    for tab in (tab_a, tab_b):
        assert tab.controller.doc.operation_log == []
        assert not tab.controller.is_dirty
        with pikepdf.Pdf.open(tab.controller.doc.working_path) as pdf:
            assert int(pdf.pages[0].get("/Rotate", 0)) == 0
    _force_close(window)


# --- crash recovery (most recently active tab only) --------------------------


def _crashed_window_with_two_dirty_tabs(tmp_path: Path) -> tuple[MainWindow, Path, Path]:
    """Two edited tabs, then the window is abandoned without ever
    closing a session - i.e. exactly what a crash leaves behind (a
    clean exit would have discarded every journal)."""
    from core.ops.organize import RotatePagesOperation

    a = _make_pdf(tmp_path / "a.pdf", 1)
    b = _make_pdf(tmp_path / "b.pdf", 1)
    window = MainWindow()
    tab_a = _open_tab(window, a)
    window.controller.apply_operation(RotatePagesOperation(angle=90))
    tab_b = _open_tab(window, b)
    window.controller.apply_operation(RotatePagesOperation(angle=180))
    # Back to A, so A - not the most recently *opened* tab - is the
    # most recently active one.
    window.tab_widget.setCurrentWidget(tab_a)
    assert tab_b.controller.is_dirty
    return window, a, b


def test_autosave_restores_only_the_most_recently_active_tab(
    qapp: QApplication, tmp_path: Path
) -> None:
    crashed, a, _b = _crashed_window_with_two_dirty_tabs(tmp_path)

    window = MainWindow()
    with patch.object(QMessageBox, "question", return_value=QMessageBox.StandardButton.Yes):
        restored = window.restore_autosaved_session()

    assert restored
    assert window.tab_widget.count() == 1  # one tab, not both
    tab = window.current_tab
    assert tab.controller.doc.source_path == a
    # Restored *unsaved* state (the 90-degree rotation that was never
    # written back to a.pdf), not a re-open of the original file.
    assert tab.controller.is_dirty
    with pikepdf.Pdf.open(tab.controller.doc.working_path) as pdf:
        assert int(pdf.pages[0].get("/Rotate", 0)) == 90
    with pikepdf.Pdf.open(a) as pdf:
        assert int(pdf.pages[0].get("/Rotate", 0)) == 0

    _force_close(window)
    _force_close(crashed)


def test_a_restored_session_is_not_offered_again_next_launch(
    qapp: QApplication, tmp_path: Path
) -> None:
    from core.session.autosave import recover_active_session

    crashed, _a, _b = _crashed_window_with_two_dirty_tabs(tmp_path)

    window = MainWindow()
    with patch.object(QMessageBox, "question", return_value=QMessageBox.StandardButton.Yes):
        window.restore_autosaved_session()
    _force_close(window)

    next_launch = MainWindow()
    with patch.object(QMessageBox, "question") as mock_question:
        assert not next_launch.restore_autosaved_session()
    mock_question.assert_not_called()
    assert recover_active_session() is None
    next_launch.close()
    _force_close(crashed)


def test_declining_recovery_opens_nothing_and_discards_the_journal(
    qapp: QApplication, tmp_path: Path
) -> None:
    from core.session.autosave import recover_active_session

    crashed, _a, _b = _crashed_window_with_two_dirty_tabs(tmp_path)

    window = MainWindow()
    with patch.object(QMessageBox, "question", return_value=QMessageBox.StandardButton.No):
        assert not window.restore_autosaved_session()

    assert window.tab_widget.count() == 0
    assert window.stack.currentWidget() is window.empty_state
    assert recover_active_session() is None
    window.close()
    _force_close(crashed)


def test_a_clean_shutdown_leaves_nothing_to_recover(qapp: QApplication, tmp_path: Path) -> None:
    from core.ops.organize import RotatePagesOperation

    src = _make_pdf(tmp_path / "src.pdf", 1)
    window = MainWindow()
    _open_tab(window, src)
    window.controller.apply_operation(RotatePagesOperation(angle=90))
    window.controller.save_as(tmp_path / "out.pdf")

    event = QCloseEvent()
    window.closeEvent(event)
    assert event.isAccepted()

    next_launch = MainWindow()
    with patch.object(QMessageBox, "question") as mock_question:
        assert not next_launch.restore_autosaved_session()
    mock_question.assert_not_called()
    next_launch.close()


# --- File > Properties... ---------------------------------------------------
#
# PropertiesDialog is a plain Python QDialog subclass (not a
# BaseToolDialog, and deliberately not a QMessageBox), so
# patch.object(PropertiesDialog, "exec", ...) genuinely intercepts -
# see CLAUDE.md on compiled .exec methods.


def test_properties_action_is_disabled_until_a_document_is_open(
    qapp: QApplication, tmp_path: Path
) -> None:
    src = _make_pdf(tmp_path / "src.pdf", 2)
    window = MainWindow()
    assert not window.properties_action.isEnabled()

    _open_tab(window, src)
    assert window.properties_action.isEnabled()

    window._close_document()
    assert not window.properties_action.isEnabled()
    window.close()


def test_properties_action_with_no_document_does_nothing_rather_than_crashing(
    qapp: QApplication,
) -> None:
    window = MainWindow()
    with patch.object(PropertiesDialog, "exec") as mock_exec:
        window.properties_action.trigger()
    mock_exec.assert_not_called()
    window.close()


def test_properties_lives_in_the_file_menu_with_the_acrobat_shortcut(
    qapp: QApplication,
) -> None:
    window = MainWindow()
    # findChildren, not `action.menu()` off the menu bar: in PySide6
    # 6.11 each `QAction.menu()` call hands Python a fresh owning
    # wrapper, and releasing it destroys the real menu - the next touch
    # then raises "Internal C++ object (QMenu) already deleted". Hit
    # for real while writing this test, not a precaution.
    file_menu = next(
        menu for menu in window.findChildren(QMenu) if menu.title().replace("&", "") == "File"
    )
    actions = file_menu.actions()
    assert window.properties_action in actions
    assert window.properties_action.shortcut() == QKeySequence("Ctrl+D")

    # Straight after Save As, and fenced off by separators from the
    # tab-management group below it.
    index = actions.index(window.properties_action)
    assert actions[index - 1].isSeparator()
    assert actions[index + 1].isSeparator()
    save_index = actions.index(window.save_as_action)
    assert save_index == index - 2
    window.close()


def test_properties_reports_the_open_documents_real_geometry_and_metadata(
    qapp: QApplication, tmp_path: Path
) -> None:
    src = _make_pdf(tmp_path / "src.pdf", 3)
    with pikepdf.Pdf.open(src, allow_overwriting_input=True) as pdf:
        pdf.docinfo["/Title"] = "Quarterly Report"
        pdf.save(src)

    window = MainWindow()
    tab = _open_tab(window, src)

    info = window._read_properties(tab)
    assert info.metadata is not None
    assert info.metadata.title == "Quarterly Report"
    assert info.geometry is not None
    assert info.geometry.page_count == 3
    assert info.file.path == src
    assert info.file.has_unsaved_changes is False
    window.close()


def test_properties_describes_the_working_copy_not_the_untouched_original(
    qapp: QApplication, tmp_path: Path
) -> None:
    """The whole point of the "unsaved changes" row: after deleting a
    page, the report must show 2 pages (the in-memory edit state) while
    still pointing at the 3-page file on disk, and say so."""
    from core.ops.organize import DeletePagesOperation

    src = _make_pdf(tmp_path / "src.pdf", 3)
    window = MainWindow()
    tab = _open_tab(window, src)
    tab.controller.apply_operation(DeletePagesOperation(pages=[3]))

    info = window._read_properties(tab)
    assert info.geometry is not None
    assert info.geometry.page_count == 2
    assert info.file.has_unsaved_changes is True
    with pikepdf.Pdf.open(src) as pdf:
        assert len(pdf.pages) == 3  # the original really is untouched

    dialog = PropertiesDialog(info)
    assert "Unsaved changes" in dialog.report_text()
    assert "3" not in dialog.report_text().split("Pages:")[1].splitlines()[0]
    _force_close(window)


def test_properties_copy_to_clipboard_puts_the_whole_report_on_the_clipboard(
    qapp: QApplication, tmp_path: Path
) -> None:
    src = _make_pdf(tmp_path / "src.pdf", 2)
    with pikepdf.Pdf.open(src, allow_overwriting_input=True) as pdf:
        pdf.docinfo["/Author"] = "A. Author"
        pdf.save(src)

    window = MainWindow()
    tab = _open_tab(window, src)
    dialog = PropertiesDialog(window._read_properties(tab))

    QGuiApplication.clipboard().clear()
    dialog.button_for_copy().click()  # the real button, so the real handler runs
    copied = QGuiApplication.clipboard().text()

    assert copied == dialog.report_text()
    # A labelled report, not a JSON/repr dump.
    assert "Document metadata" in copied
    assert "Author:" in copied
    assert "A. Author" in copied
    assert "Pages:" in copied
    assert str(src) in copied
    assert not copied.lstrip().startswith(("{", "["))
    window.close()


def test_properties_edit_metadata_goes_through_the_real_undoable_tool_path(
    qapp: QApplication, tmp_path: Path
) -> None:
    """Editing from the Properties dialog must land in the undo stack
    and the audit log exactly as Tools > Metadata does - it runs the
    same _run_tool path, not a second hand-rolled apply route - and the
    dialog must then refresh instead of showing the old values."""
    src = _make_pdf(tmp_path / "src.pdf", 1)
    window = MainWindow()
    tab = _open_tab(window, src)
    audit_entries_before = len(window.audit_log.read_all())

    dialog = PropertiesDialog(
        window._read_properties(tab),
        lambda: window._edit_metadata_from_properties(tab),
    )
    assert dialog.edit_metadata_button.isEnabled()
    assert "(not set)" in dialog.report_text()

    def fake_metadata(self: MetadataDialog) -> QDialog.DialogCode:
        self.title.setText("Edited From Properties")
        self.author.setText("Radwan")
        return QDialog.DialogCode.Accepted

    with patch.object(MetadataDialog, "exec", fake_metadata):
        dialog.edit_metadata_button.click()

    # The edit really happened, through the ordinary Operation path.
    applied = tab.controller.doc.operation_log[-1].serialize()
    assert applied["type"] == "set_metadata"
    assert window.undo_action.isEnabled()
    assert len(window.audit_log.read_all()) == audit_entries_before + 1
    with pikepdf.Pdf.open(tab.controller.doc.working_path) as pdf:
        assert str(pdf.docinfo["/Title"]) == "Edited From Properties"

    # ...and the still-open report refreshed rather than going stale.
    refreshed = dialog.report_text()
    assert "Edited From Properties" in refreshed
    assert "Radwan" in refreshed
    _force_close(window)


def test_properties_edit_metadata_cancelled_leaves_the_report_untouched(
    qapp: QApplication, tmp_path: Path
) -> None:
    src = _make_pdf(tmp_path / "src.pdf", 1)
    window = MainWindow()
    tab = _open_tab(window, src)
    dialog = PropertiesDialog(
        window._read_properties(tab),
        lambda: window._edit_metadata_from_properties(tab),
    )
    before = dialog.report_text()

    def fake_cancel(self: MetadataDialog) -> QDialog.DialogCode:
        return QDialog.DialogCode.Rejected

    with patch.object(MetadataDialog, "exec", fake_cancel):
        dialog.edit_metadata_button.click()

    assert dialog.report_text() == before
    assert tab.controller.doc.operation_log == []
    window.close()


def test_properties_dialog_opens_from_the_menu_action(
    qapp: QApplication, tmp_path: Path
) -> None:
    src = _make_pdf(tmp_path / "src.pdf", 1)
    window = MainWindow()
    _open_tab(window, src)

    shown: list[PropertiesDialog] = []

    def fake_exec(self: PropertiesDialog) -> QDialog.DialogCode:
        shown.append(self)
        return QDialog.DialogCode.Accepted

    with patch.object(PropertiesDialog, "exec", fake_exec):
        window.properties_action.trigger()

    assert len(shown) == 1
    assert "src.pdf" in shown[0].report_text()
    window.close()


def test_properties_of_a_password_protected_document_degrades_without_crashing(
    qapp: QApplication, tmp_path: Path
) -> None:
    """A real, reachable state: Protect encrypts the working copy with
    a user password, so the inspector cannot reopen it."""
    from core.ops.security import ProtectOperation

    src = _make_pdf(tmp_path / "src.pdf", 1)
    window = MainWindow()
    tab = _open_tab(window, src)
    tab.controller.apply_operation(ProtectOperation(user_password="hunter2"))

    info = window._read_properties(tab)
    assert info.password_protected is True
    dialog = PropertiesDialog(info)
    report = dialog.report_text()
    assert "Encrypted:" in report
    assert "Unavailable" in report
    assert str(src) in report  # the on-disk facts still work
    _force_close(window)


def test_properties_of_a_document_with_no_file_on_disk(
    qapp: QApplication, tmp_path: Path
) -> None:
    """Merge builds a document from scratch, so there is no source file
    to stat - the report says so rather than showing a blank path."""
    inputs = [_make_pdf(tmp_path / f"in{i}.pdf", 1) for i in range(2)]
    window = MainWindow()

    def fake_merge(self: MergeDialog) -> QDialog.DialogCode:
        for path in inputs:
            self.file_list.addItem(str(path))
        return QDialog.DialogCode.Accepted

    with patch.object(MergeDialog, "exec", fake_merge):
        window._run_tool("merge", MergeDialog)

    tab = window.current_tab
    assert tab is not None
    info = window._read_properties(tab)
    assert info.file.path is None
    assert info.geometry is not None
    assert info.geometry.page_count == 2
    assert "not saved to disk yet" in PropertiesDialog(info).report_text()
    _force_close(window)


@pytest.mark.skipif(
    sys.platform != "linux", reason="/proc/self/fd is the Linux way to see open handles"
)
def test_properties_holds_no_handle_on_the_working_file(
    qapp: QApplication, tmp_path: Path
) -> None:
    """The inspector must not keep the session working copy open.

    CLAUDE.md records a real Windows CI failure (WinError 32) where a
    leaked QPdfDocument handle turned "securely wipe the confidential
    working copy on close" into a SecurityError. This dialog reads via
    a `with pikepdf.Pdf.open(...)` block and holds only plain data, so
    it should never take a handle at all - checked against the OS
    rather than inferred, because the wipe below succeeds on Linux
    whether a handle leaked or not.
    """
    src = _make_pdf(tmp_path / "src.pdf", 1)
    window = MainWindow()
    tab = _open_tab(window, src)
    working = tab.controller.doc.working_path

    def open_files() -> set[str]:
        found = set()
        for fd in os.listdir("/proc/self/fd"):
            with contextlib.suppress(OSError):
                found.add(os.readlink(f"/proc/self/fd/{fd}"))
        return found

    def fake_exec(self: PropertiesDialog) -> QDialog.DialogCode:
        assert str(working) not in open_files(), "properties dialog is holding the working file"
        return QDialog.DialogCode.Accepted

    with patch.object(PropertiesDialog, "exec", fake_exec):
        window.properties_action.trigger()

    assert str(working) not in open_files()
    session_dir = working.parent
    tab.controller.close_session()
    assert not session_dir.exists()
    window.close()


# --- PDF -> external-format exports -----------------------------------------
#
# These five conversions used to be applied to the tab like any other
# operation, which replaced its PDF working file with a .docx/.pptx/
# .xlsx/.html/.jpg. The thumbnail grid cannot render one, so the window
# went blank and the document appeared to vanish - and the converted
# file, left in the private session dir, was securely wiped when the
# tab closed. They are exports now: the file goes where the user asks,
# and the open document is untouched.


_EXPORT_CASES = [
    ("pdf_to_docx", {}, ".docx"),
    ("pdf_to_pptx", {}, ".pptx"),
    ("pdf_to_xlsx", {}, ".xlsx"),
    ("pdf_to_html", {}, ".html"),
    ("pdf_to_jpg", {"page": 1}, ".jpg"),
]


@pytest.mark.parametrize(("tool_id", "values", "suffix"), _EXPORT_CASES)
def test_converting_from_pdf_writes_a_file_and_leaves_the_document_open(
    qapp: QApplication,
    tmp_path: Path,
    tool_id: str,
    values: dict[str, Any],
    suffix: str,
) -> None:
    src = _make_pdf(tmp_path / "src.pdf", 2)
    destination = tmp_path / f"exported{suffix}"
    window = MainWindow()
    tab = _open_tab(window, src)
    dialog_cls = TOOL_DIALOGS[tool_id]

    with (
        patch.object(dialog_cls, "exec", lambda self: QDialog.DialogCode.Accepted),
        patch.object(dialog_cls, "values", lambda self: values),
        patch("gui.main_window.QFileDialog.getSaveFileName", return_value=(str(destination), "")),
    ):
        window._run_tool(tool_id, dialog_cls)

    # The converted file is where the user asked for it, and is real.
    assert destination.exists()
    assert destination.stat().st_size > 0
    # The window still shows the PDF - this is the blank-screen symptom.
    assert tab.thumbnail_list.count() == 2
    working_path = tab.controller.doc.working_path
    assert working_path is not None
    assert working_path.suffix == ".pdf"
    # An export is not an edit: nothing to undo, nothing unsaved.
    assert tab.controller.doc.operation_log == []
    assert not tab.controller.is_dirty
    assert not window.undo_action.isEnabled()
    _force_close(window)


def test_an_export_survives_the_session_being_wiped(qapp: QApplication, tmp_path: Path) -> None:
    """The other half of the bug: the converted file used to live only
    in the private session dir, which is securely wiped when the tab
    closes - so the user's Word file was destroyed rather than
    delivered. It has to outlive the session that produced it."""
    src = _make_pdf(tmp_path / "src.pdf", 1)
    destination = tmp_path / "exported.docx"
    window = MainWindow()
    tab = _open_tab(window, src)
    working_path = tab.controller.doc.working_path
    assert working_path is not None
    session_dir = working_path.parent

    with (
        patch.object(PdfToDocxDialog, "exec", lambda self: QDialog.DialogCode.Accepted),
        patch("gui.main_window.QFileDialog.getSaveFileName", return_value=(str(destination), "")),
    ):
        window._run_tool("pdf_to_docx", PdfToDocxDialog)

    _force_close(window)

    assert not session_dir.exists(), "the tab's session dir should have been wiped"
    assert destination.exists(), "the exported file was wiped along with the session"
    assert destination.stat().st_size > 0


def test_cancelling_the_export_destination_changes_nothing(
    qapp: QApplication, tmp_path: Path
) -> None:
    src = _make_pdf(tmp_path / "src.pdf", 2)
    window = MainWindow()
    tab = _open_tab(window, src)

    with (
        patch.object(PdfToDocxDialog, "exec", lambda self: QDialog.DialogCode.Accepted),
        patch("gui.main_window.QFileDialog.getSaveFileName", return_value=("", "")),
    ):
        window._run_tool("pdf_to_docx", PdfToDocxDialog)

    assert tab.thumbnail_list.count() == 2
    assert tab.controller.doc.operation_log == []
    assert not tab.controller.is_dirty
    _force_close(window)


def test_an_export_is_recorded_in_the_audit_log(qapp: QApplication, tmp_path: Path) -> None:
    """Every other path that applies an Operation records to the audit
    trail, and an export writes a confidential document out to a new
    file - exactly what the trail exists for."""
    src = _make_pdf(tmp_path / "src.pdf", 1)
    destination = tmp_path / "exported.docx"
    window = MainWindow()
    _open_tab(window, src)
    entries_before = len(window.audit_log.read_all())

    with (
        patch.object(PdfToDocxDialog, "exec", lambda self: QDialog.DialogCode.Accepted),
        patch("gui.main_window.QFileDialog.getSaveFileName", return_value=(str(destination), "")),
    ):
        window._run_tool("pdf_to_docx", PdfToDocxDialog)

    entries = window.audit_log.read_all()
    assert len(entries) == entries_before + 1
    # The field itself, not a substring of the whole entry's repr: on
    # Windows a path's backslashes come back doubled through JSON, so
    # `str(destination) in str(entry)` is false for a correctly
    # recorded entry (caught by CI on Windows, not locally).
    assert entries[-1]["document"] == str(destination)
    assert entries[-1]["operation"]["type"] == "pdf_to_docx"
    _force_close(window)


def test_an_export_extension_is_added_only_when_the_user_typed_none(
    qapp: QApplication, tmp_path: Path
) -> None:
    src = _make_pdf(tmp_path / "src.pdf", 1)
    window = MainWindow()
    _open_tab(window, src)
    bare = tmp_path / "no_extension"

    with (
        patch.object(PdfToDocxDialog, "exec", lambda self: QDialog.DialogCode.Accepted),
        patch("gui.main_window.QFileDialog.getSaveFileName", return_value=(str(bare), "")),
    ):
        window._run_tool("pdf_to_docx", PdfToDocxDialog)

    assert (tmp_path / "no_extension.docx").exists()
    assert not bare.exists()

    # An extension the user chose deliberately is left alone.
    chosen = tmp_path / "report.doc"
    with (
        patch.object(PdfToDocxDialog, "exec", lambda self: QDialog.DialogCode.Accepted),
        patch("gui.main_window.QFileDialog.getSaveFileName", return_value=(str(chosen), "")),
    ):
        window._run_tool("pdf_to_docx", PdfToDocxDialog)

    assert chosen.exists()
    _force_close(window)


def test_a_failing_export_reports_the_error_and_leaves_the_document_open(
    qapp: QApplication, tmp_path: Path
) -> None:
    src = _make_pdf(tmp_path / "src.pdf", 2)
    destination = tmp_path / "out.jpg"
    window = MainWindow()
    tab = _open_tab(window, src)
    errors: list[str] = []

    with (
        patch.object(PdfToJpgDialog, "exec", lambda self: QDialog.DialogCode.Accepted),
        # Page 99 of a 2-page document: rejected by the operation itself.
        patch.object(PdfToJpgDialog, "values", lambda self: {"page": 99}),
        patch("gui.main_window.QFileDialog.getSaveFileName", return_value=(str(destination), "")),
        patch.object(MainWindow, "_show_error", lambda self, exc: errors.append(str(exc))),
    ):
        window._run_tool("pdf_to_jpg", PdfToJpgDialog)

    assert errors, "a failed export must report the error"
    assert not destination.exists()
    assert tab.thumbnail_list.count() == 2
    assert tab.controller.doc.operation_log == []
    _force_close(window)
